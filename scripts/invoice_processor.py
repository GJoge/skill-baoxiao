#!/usr/bin/env python3
"""
报销发票处理脚本
功能：识别发票类型、提取数据、校验合理性、填写Excel
"""

import pdfplumber
import pdf2image
import pytesseract
import os
import re
import json
import shutil
import subprocess
import glob
import sys
from datetime import datetime
from openpyxl import load_workbook
from collections import defaultdict

# PDF写入相关
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from pypdf import PdfWriter, PdfReader

# ============ 配置常量 ============

# 发票类型关键字
TYPE_KEYWORDS = {
    'tuigai': ['退票费', '改签费', '退票', '改签', '变更'],
    'huoche': ['铁路', '车次', '二等座', '高铁', '动车', '列车', '12306', '出发站', '到达站', '铁路电子客票'],
    'jipiao': ['航空', '航班', '承运人', '民航', '座位等级', '国内国际标识', '机场', '登机', '客票级别', '航班号'],
    'zhusu': ['住宿', '酒店', '宾馆', '住宿费', '房费', '旅店', '客房', '住宿服务'],
    'didi': ['滴滴', '网约车', '行程单', '出行', '快车', '专车'],
    'toll': ['通行费', '入口站', '出口站', '入口时间', '出口时间', '收费公路通行费', '高速公路通行费', 'ETC'],
}

# 其他交通类发票关键字（归入滴滴类别）
TRANSPORT_KEYWORDS = [
    '地铁', '轨道交通', '公交', '公共汽车', '出租车', '的士', '一卡通',
    '城市通', '交通卡', '乘车码', '扫码乘车', '公共交通'
]

# 高速通行费强关键字。需优先于机票/滴滴判断，避免“机场高速”误判为机票。
TOLL_KEYWORDS = [
    '通行费', '入口站', '出口站', '入口时间', '出口时间',
    '收费公路通行费', '高速公路通行费', '代收通行费',
    '通行费电子行程单', 'ETC', '广东联合电子服务股份有限公司'
]

# 用户确认的发票类型映射（用于记住用户选择）
USER_CONFIRMED_TYPES = {}

# 金额合理范围（用于校验）
AMOUNT_RANGES = {
    'tuigai': (5, 2000),
    'huoche': (20, 3000),
    'jipiao': (300, 8000),
    'zhusu': (100, 3000),
    'didi': (2, 800),  # 包含地铁、公交等小额交通费用（通常2元起）
    'toll': (1, 1000),
}

# OCR配置
OCR_DPI_LOW = 200   # 用于类型识别
OCR_DPI_HIGH = 400  # 用于日期精确识别


# ============ 工作区准备函数 ============

def prepare_invoice_workspace(invoice_dir='发票', backup_dir='bak'):
    """
    准备发票工作区。

    - 如果 bak 不存在：将当前发票目录复制为 bak，保留原始副本。
    - 如果 bak 已存在：删除当前发票目录，并从 bak 复制出干净的发票目录。
    """
    if os.path.exists(backup_dir):
        if not os.path.isdir(backup_dir):
            raise NotADirectoryError(f"备份路径不是目录: {backup_dir}")
        if os.path.exists(invoice_dir):
            shutil.rmtree(invoice_dir)
        shutil.copytree(backup_dir, invoice_dir)
        print(f"✓ 已从备份恢复干净发票目录: {backup_dir} -> {invoice_dir}")
        return "restored_from_backup"

    if not os.path.exists(invoice_dir):
        raise FileNotFoundError(f"发票目录不存在，无法创建备份: {invoice_dir}")
    if not os.path.isdir(invoice_dir):
        raise NotADirectoryError(f"发票路径不是目录: {invoice_dir}")

    shutil.copytree(invoice_dir, backup_dir)
    print(f"✓ 已创建发票备份: {invoice_dir} -> {backup_dir}")
    return "backup_created"


# ============ 配置文件加载函数 ============

def load_config(config_path='config.yaml'):
    """
    加载配置文件
    优先查找用户指定的路径，然后查找脚本所在目录
    返回: 配置字典，如果配置文件不存在则返回空字典
    """
    import yaml

    # 尝试多个位置查找配置文件
    possible_paths = [
        config_path,  # 当前目录或指定路径
        os.path.join(os.path.dirname(__file__), '..', config_path),  # 脚本上级目录
        os.path.join(os.path.dirname(__file__), config_path),  # 脚本同级目录
        os.path.expanduser(f'~/.config/baoxiao/{config_path}'),  # 用户配置目录
    ]

    for path in possible_paths:
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                    print(f"✓ 已加载配置文件: {path}")
                    return config or {}
            except Exception as e:
                print(f"⚠️ 配置文件加载失败 {path}: {e}")
                continue

    return {}

def get_city_unit_map(config=None, cli_map=None):
    """
    获取城市到单位的映射
    优先级: 命令行参数 > 配置文件 > 空字典

    Args:
        config: 配置字典（从load_config获取）
        cli_map: 命令行传入的映射字典

    Returns:
        城市到单位的映射字典
    """
    # 从配置文件获取
    config_map = {}
    if config and 'city_units' in config:
        config_map = config['city_units']

    # 合并：命令行参数优先级最高
    if cli_map:
        merged = config_map.copy()
        merged.update(cli_map)
        return merged

    return config_map


# ============ 微信支付账单处理函数 ============

def find_wechat_bill_file(input_dir):
    """
    在发票文件夹中查找包含'微信支付账单'的xlsx文件
    排除临时文件（以.~或~$开头）
    返回: 文件路径或None
    """
    if not os.path.exists(input_dir):
        return None

    for filename in os.listdir(input_dir):
        # 跳过临时文件
        if filename.startswith('.~') or filename.startswith('~$'):
            continue
        if '微信支付账单' in filename and filename.endswith('.xlsx'):
            return os.path.join(input_dir, filename)
    return None


def load_wechat_bill_data(bill_path):
    """
    读取微信支付账单xlsx文件，提取金额(元)列(F列)、交易单号列(I列)和交易对方列(C列)
    对于包含"已退款(¥xx.xx)"的记录，计算退票费（原金额 - 已退款金额）并单独存储
    返回: {金额: [交易单号列表]} 的字典、金额到交易对方的映射、退票费映射
    """
    try:
        wb = load_workbook(bill_path, read_only=True)
        ws = wb.active

        amount_to_trans = {}  # {金额: [交易单号列表]} - 支持相同金额
        amount_to_merchant = {}  # 金额到交易对方的映射
        refund_fee_to_trans = {}  # {退票费: [交易单号列表]} - 用于匹配tuigai单据

        # 从第2行开始读取（跳过表头）
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            # C列是第3列（索引2），F列是第6列（索引5），H列是第8列（索引7），I列是第9列（索引8）
            if len(row) >= 9:
                merchant = row[2]  # C列 - 交易对方
                amount = row[5]  # F列 - 金额(元)
                status = row[7]  # H列 - 当前状态
                trans_no = row[8]  # I列 - 交易单号

                # 删除已全额退款的记录（不是真实消费，不应匹配交易单号）
                if status and '已全额退款' in str(status):
                    continue  # 跳过这些记录，不添加到amount_to_trans

                # 处理部分退款 - 计算退票费
                refunded_amount = None
                refund_fee = None
                if status and '已退款' in str(status) and '(' in str(status) and ')' in str(status):
                    # 提取括号中的已退款金额，例如 "已退款(¥100.00)"
                    try:
                        import re
                        match = re.search(r'[¥￥](\d+\.?\d*)', str(status))
                        if match:
                            refunded_amount = float(match.group(1))
                    except:
                        pass

                if amount is not None and trans_no is not None:
                    # 处理金额（可能是字符串或数字）
                    if isinstance(amount, str):
                        # 去掉正负号和货币符号，取绝对值
                        amount_str = amount.replace('¥', '').replace('￥', '').replace(',', '').strip()
                        try:
                            amount_val = abs(float(amount_str))
                        except:
                            continue
                    else:
                        amount_val = abs(float(amount))

                    # 如果存在部分退款，计算退票费并单独存储
                    if refunded_amount is not None and amount_val > 0:
                        refund_fee = amount_val - refunded_amount
                        if refund_fee > 0:
                            # 存储退票费用于匹配tuigai单据
                            if refund_fee not in refund_fee_to_trans:
                                refund_fee_to_trans[refund_fee] = []
                            refund_fee_to_trans[refund_fee].append({
                                'trans_no': str(trans_no).strip(),
                                'original_amount': amount_val,
                                'refunded_amount': refunded_amount,
                                'merchant': str(merchant).strip() if merchant else ''
                            })
                            # 仍然存储原金额用于普通匹配
                            if amount_val not in amount_to_trans:
                                amount_to_trans[amount_val] = []
                            amount_to_trans[amount_val].append(str(trans_no).strip())
                            if merchant:
                                amount_to_merchant[amount_val] = str(merchant).strip()
                            continue

                    # 只保留非零金额
                    if amount_val > 0:
                        # 使用列表存储相同金额的交易单号，解决重复金额覆盖问题
                        if amount_val not in amount_to_trans:
                            amount_to_trans[amount_val] = []
                        amount_to_trans[amount_val].append(str(trans_no).strip())
                        if merchant:
                            amount_to_merchant[amount_val] = str(merchant).strip()

        wb.close()
        return amount_to_trans, amount_to_merchant, refund_fee_to_trans

    except Exception as e:
        print(f"  ⚠️ 读取微信支付账单失败: {e}")
        return {}, {}, {}


def extract_amounts_from_didi_trip(pdf_path):
    """
    从滴滴出行行程报销单中提取所有行程金额（多笔行程）
    返回: [金额列表] - 每笔行程一个金额
    """
    amounts = []
    try:
        # 提取PDF文本
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        # 如果文本层为空，使用OCR
        if len(text.strip()) < 50:
            images = pdf2image.convert_from_path(pdf_path, dpi=OCR_DPI_HIGH)
            text = ""
            for img in images:
                text += pytesseract.image_to_string(img, lang='chi_sim')

        # 滴滴行程单表格格式：
        # 序号 车型 上车时间 城市 起点 终点 里程[公里] 金额[元] 备注
        # 1 快车 02-01 18:14 北京市 ... 3.1 12.41
        # 金额通常在行尾，是xx.xx格式的数字

        # 方法1: 匹配行程行中的金额（在里程数后面的金额）
        # 匹配包含时间格式的行，提取其中的金额数字
        lines = text.split('\n')
        for line in lines:
            # 跳过表头行
            if '金额[元]' in line or '金额' in line and '里程' in line:
                continue
            # 匹配包含时间格式(xx:xx)的行
            if re.search(r'\d{1,2}:\d{2}', line):
                # 提取行中所有xx.xx格式的数字
                matches = re.findall(r'(\d+\.\d{2})', line)
                for match in matches:
                    try:
                        val = float(match)
                        # 筛选合理范围内的金额（5-500元），排除里程数（通常是x.x格式）
                        if 5 <= val <= 500 and val not in amounts:
                            amounts.append(val)
                    except:
                        continue

        # 方法2: 如果方法1没有提取到，尝试匹配所有合理的金额数字
        if not amounts:
            # 查找所有xx.xx格式的数字
            all_matches = re.findall(r'\b(\d+\.\d{2})\b', text)
            for match in all_matches:
                try:
                    val = float(match)
                    # 筛选合理范围内的金额（5-500元）
                    if 5 <= val <= 500 and val not in amounts:
                        amounts.append(val)
                except:
                    continue

        # 方法3: 匹配¥符号或"元"字后面的金额
        if not amounts:
            patterns = [
                r'[¥￥]\s*(\d+\.\d{2})',  # ¥12.34
                r'(\d+\.\d{2})\s*元',     # 12.34元
            ]
            for pattern in patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    try:
                        val = float(match)
                        if 5 <= val <= 500 and val not in amounts:
                            amounts.append(val)
                    except:
                        continue

    except Exception as e:
        print(f"  ⚠️ 提取滴滴行程金额失败 {pdf_path}: {e}")

    return amounts


def add_transaction_numbers_to_pdf(input_pdf, output_pdf, trans_numbers):
    """
    将交易单号添加到PDF最后一页的底部居中位置
    trans_numbers: [(编号, 交易单号), ...]
    """
    try:
        # 读取原始PDF
        reader = PdfReader(input_pdf)

        # 获取最后一页的尺寸
        last_page = reader.pages[-1]
        page_width = float(last_page.mediabox.width)
        page_height = float(last_page.mediabox.height)

        # 创建一个新的PDF页面，包含交易单号文本
        from io import BytesIO
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))

        # 尝试使用系统中文字体
        font_name = "Helvetica"
        try:
            # 尝试使用常见的中文字体路径
            font_paths = [
                '/usr/share/fonts/custom/msyh.ttc',  # 微软雅黑
                '/usr/share/fonts/custom/SIMSUN.ttf',  # 宋体
                '/usr/share/fonts/custom/simkai.ttf',  # 楷体
                '/usr/share/fonts/custom/simhei.ttf',  # 黑体
                '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',  # Droid Fallback
                '/usr/local/share/fonts/custom/msyh.ttc',  # 微软雅黑(旧路径)
                '/usr/local/share/fonts/custom/SIMSUN.ttf',  # 宋体(旧路径)
                '/usr/local/share/fonts/custom/simkai.ttf',  # 楷体(旧路径)
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            ]
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('CustomFont', font_path))
                    font_name = 'CustomFont'
                    break
        except Exception as e:
            print(f"  ⚠️ 字体注册失败: {e}, 使用默认字体")

        can.setFont(font_name, 10)

        # 在页面底部居中添加交易单号
        # 判断是否有多个交易单号
        has_multiple = len(trans_numbers) > 1

        # 计算起始y位置（根据交易单号数量）
        # 序号小的画在上面（y坐标大），序号大的画在下面（y坐标小）
        base_y = 30 if has_multiple else 7
        line_height = 15
        y_pos = base_y + (len(trans_numbers) - 1) * line_height

        for idx, trans_no in trans_numbers:
            # 只有多个交易单号时才显示序号
            if has_multiple:
                text = f"[{idx}] 交易单号: {trans_no}"
            else:
                text = f"交易单号: {trans_no}"
            # 计算文本宽度，实现居中
            text_width = can.stringWidth(text, font_name, 10)
            x_pos = (page_width - text_width) / 2  # 居中位置
            can.drawString(x_pos, y_pos, text)
            y_pos -= line_height  # 向上移动

        can.save()
        packet.seek(0)

        # 读取新创建的PDF页面
        new_pdf = PdfReader(packet)

        # 创建一个新的writer，合并原始页面和新的文本层
        final_writer = PdfWriter()

        # 复制除最后一页外的所有页面
        for i in range(len(reader.pages) - 1):
            final_writer.add_page(reader.pages[i])

        # 合并最后一页和文本层
        last_page = reader.pages[-1]
        text_page = new_pdf.pages[0]

        # 将文本页合并到最后一页
        last_page.merge_page(text_page)
        final_writer.add_page(last_page)

        # 保存结果
        with open(output_pdf, 'wb') as f:
            final_writer.write(f)

        return True

    except Exception as e:
        print(f"  ⚠️ 添加交易单号失败: {e}")
        # 如果失败，复制原文件
        shutil.copy2(input_pdf, output_pdf)
        return False


def match_and_add_transaction_numbers(input_dir, work_dir, report_path='data_report.json', force_mark=False):
    """
    主函数：查找微信支付账单，匹配发票金额，添加交易单号到PDF
    使用data_report中的file_details数据，避免重复提取金额

    Args:
        input_dir: 发票目录
        work_dir: 工作目录
        report_path: 报告文件路径
        force_mark: 是否强制重新标记（忽略已有标记）

    返回: 处理是否成功的布尔值
    """
    print("\n[步骤0] 检查微信支付账单...")

    # 1. 查找微信支付账单文件
    bill_path = find_wechat_bill_file(input_dir)
    if not bill_path:
        print("  ℹ️ 未找到微信支付账单文件，跳过此步骤")
        return False

    print(f"  ✓ 找到微信支付账单: {os.path.basename(bill_path)}")

    # 2. 加载账单数据（包括退票费映射用于tuigai单据）
    amount_to_trans, amount_to_merchant, refund_fee_to_trans = load_wechat_bill_data(bill_path)
    if not amount_to_trans and not refund_fee_to_trans:
        print("  ⚠️ 未能从账单中提取有效数据")
        return False

    print(f"  ✓ 从账单中提取了 {len(amount_to_trans)} 条交易记录")
    if refund_fee_to_trans:
        print(f"  ✓ 从账单中提取了 {sum(len(v) for v in refund_fee_to_trans.values())} 条退票费记录")

    # 3. 加载data_report获取文件明细
    file_details_map = {}
    if os.path.exists(report_path):
        try:
            with open(report_path, 'r', encoding='utf-8') as f:
                report = json.load(f)
                for fd in report.get('file_details', []):
                    file_details_map[fd.get('filename', '')] = fd
            print(f"  ✓ 从{report_path}加载了 {len(file_details_map)} 个文件明细")
        except Exception as e:
            print(f"  ⚠️ 加载{report_path}失败: {e}，将使用实时提取")

    # 4. 遍历发票目录中的所有PDF文件
    pdf_files = [f for f in os.listdir(input_dir) if f.endswith('.pdf')]
    if not pdf_files:
        print("  ⚠️ 发票目录中没有PDF文件")
        return False

    def get_file_type(pdf_file):
        file_info = file_details_map.get(pdf_file, {})
        return file_info.get('type', '')

    def match_sort_key(pdf_file):
        ftype = get_file_type(pdf_file)
        if ftype == 'didi_trip' or '滴滴出行行程报销单' in pdf_file:
            return (0, pdf_file)
        return (1, pdf_file)

    def pdf_has_transaction_number(pdf_path):
        try:
            from pypdf import PdfReader
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return "交易单号" in text
        except:
            return False

    def take_first_transaction(amount):
        if amount in amount_to_trans and amount_to_trans[amount]:
            return amount_to_trans[amount].pop(0)
        return None

    def take_fuzzy_didi_transaction(amount, pdf_file):
        for bill_amount, trans_numbers_list in amount_to_trans.items():
            diff = bill_amount - amount
            # 差额在0.01到10元之间，且交易对方是"滴滴出行"
            if 0.01 <= diff <= 10:
                merchant = amount_to_merchant.get(bill_amount, '')
                if '滴滴' in merchant and trans_numbers_list:
                    trans_no = trans_numbers_list.pop(0)
                    # 记录模糊匹配情况
                    fuzzy_matches.append({
                        'file': pdf_file,
                        'invoice_amount': amount,
                        'bill_amount': bill_amount,
                        'diff': diff,
                        'merchant': merchant
                    })
                    return trans_no
        return None

    def take_refund_transaction(amount):
        if amount in refund_fee_to_trans and refund_fee_to_trans[amount]:
            refund_info = refund_fee_to_trans[amount].pop(0)
            return refund_info['trans_no']
        return None

    print(f"\n[步骤0.1] 开始匹配交易单号到发票PDF...")

    matched_count = 0
    fuzzy_matches = []  # 存储模糊匹配的记录
    for pdf_file in sorted(pdf_files, key=match_sort_key):
        pdf_path = os.path.join(input_dir, pdf_file)

        # 跳过非发票PDF（如已经处理过的备份文件）
        if pdf_file.startswith('.'):
            continue

        # 从file_details获取该文件的信息
        file_info = file_details_map.get(pdf_file, {})
        ftype = file_info.get('type', '')

        # 检查PDF是否已经包含交易单号，避免重复添加（除非force_mark=True）。
        # 必须先检查，再消费账单里的交易单号。
        if not force_mark and pdf_has_transaction_number(pdf_path):
            print(f"  ✓ {pdf_file}: 已有交易单号，跳过添加")
            continue

        trans_numbers = []

        if ftype == 'didi_trip' or '滴滴出行行程报销单' in pdf_file:
            # 滴滴行程单：使用预提取的行程金额列表
            amounts = file_info.get('trip_amounts', [])
            # 如果没有预提取的数据，尝试实时提取
            if not amounts and file_info.get('amount'):
                amounts = [file_info['amount']]
            if amounts:
                idx = 1
                for amount in amounts:
                    # 第一步：精确匹配（允许0.01元误差）
                    trans_no = take_first_transaction(amount)
                    if trans_no:
                        trans_numbers.append((idx, trans_no))
                        idx += 1
                        continue  # 精确匹配成功后跳过模糊匹配，继续处理下一笔金额

                    # 第二步：如果精确匹配失败，尝试模糊匹配（仅适用于滴滴行程单）
                    trans_no = take_fuzzy_didi_transaction(amount, pdf_file)
                    if trans_no:
                        trans_numbers.append((idx, trans_no))
                        idx += 1
        else:
            # 其他发票（机票、住宿、退改签等）：使用预提取的金额
            amount = file_info.get('amount')
            # 如果没有预提取的数据，尝试实时提取
            if amount is None:
                amount, _ = extract_amount_from_pdf(pdf_path, ftype)
            if amount:
                # tuigai类型使用退票费映射匹配
                if ftype == 'tuigai':
                    trans_no = take_refund_transaction(amount)
                    if trans_no:
                        trans_numbers.append((1, trans_no))
                else:
                    # 其他类型使用普通金额映射
                    trans_no = take_first_transaction(amount)
                    if trans_no:
                        trans_numbers.append((1, trans_no))

        # 如果有匹配的交易单号，添加到PDF
        if trans_numbers:
            output_pdf = os.path.join(work_dir, f'标记_{pdf_file}')
            if add_transaction_numbers_to_pdf(pdf_path, output_pdf, trans_numbers):
                # 用标记后的文件替换原文件
                shutil.move(output_pdf, pdf_path)
                trans_list = ", ".join([f"[{n}]" for n, _ in trans_numbers])
                print(f"  ✓ {pdf_file}: 添加了 {len(trans_numbers)} 个交易单号 {trans_list}")
                matched_count += 1

    print(f"\n  ✓ 共为 {matched_count} 个PDF文件添加了交易单号")

    # 5. 如果有模糊匹配的情况，提醒用户
    if fuzzy_matches:
        print("\n  ⚠️ 以下滴滴行程单使用了模糊金额匹配（可能包含高速路桥费）：")
        for match in fuzzy_matches:
            print(f"    - {match['file']}: 发票金额¥{match['invoice_amount']} -> 账单金额¥{match['bill_amount']} (差额¥{match['diff']:.2f})")
        print("  请人工核对这些匹配是否正确。")

    return matched_count > 0


# ============ PDF处理函数 ============

def ocr_pdf(pdf_path, dpi=OCR_DPI_LOW):
    """使用OCR提取PDF中的文字"""
    try:
        images = pdf2image.convert_from_path(pdf_path, dpi=dpi)
        text = ""
        for img in images:
            text += pytesseract.image_to_string(img, lang='chi_sim+eng')
        return text
    except Exception as e:
        print(f"OCR错误 {pdf_path}: {e}")
        return ""


def extract_text_from_pdf(pdf_path):
    """从PDF提取文本，如果是图片则使用OCR"""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
    except Exception as e:
        pass

    if not text.strip():
        text = ocr_pdf(pdf_path, dpi=OCR_DPI_LOW)

    return text


# ============ 识别函数 ============

def identify_invoice_type(text, filename, interactive=True):
    """根据内容识别发票类型

    Args:
        text: 发票文本内容
        filename: 文件名（用于显示）
        interactive: 是否交互式询问用户（默认为True）

    Returns:
        发票类型字符串，或None（如果用户取消）
    """
    text_lower = text.lower()

    # 检查缓存的用户确认结果
    if filename in USER_CONFIRMED_TYPES:
        return USER_CONFIRMED_TYPES[filename]

    if is_toll_invoice_text(text):
        print(f"  检测到高速通行费关键词，归入高速费类别: {filename}")
        return 'toll'

    scores = {}
    for inv_type, keywords in TYPE_KEYWORDS.items():
        scores[inv_type] = sum(1 for k in keywords if k in text)

    # 优先检查退改签特征（包含退票费/退票手续费/改签费的PDF应优先识别为退改签）
    if any(k in text for k in ['退票费', '退票手续费', '改签费', '变更手续费', '退改签']):
        scores['tuigai'] += 20  # 提高退改签优先级
        # 退票类文本常同时带有航空关键词，直接压过机票分数
        scores['jipiao'] = min(scores['jipiao'], 1)

    max_type = max(scores, key=scores.get)
    max_score = scores[max_type]

    # 检查是否是其他交通类发票（归入滴滴类别）
    transport_score = sum(1 for k in TRANSPORT_KEYWORDS if k in text)
    if transport_score > 0 and max_type != 'tuigai':
        print(f"  检测到交通类发票关键词（地铁/公交/出租车等），归入滴滴类别: {filename}")
        return 'didi'

    if max_score == 0:
        # 无法识别，询问用户
        if interactive and sys.stdin.isatty():
            print(f"\n⚠️ 无法识别发票类型: {filename}")
            print(f"   文件内容片段: {text[:200]}...")
            print(f"\n请选择发票类型:")
            print(f"  1. 机票 (jipiao)")
            print(f"  2. 火车票 (huoche)")
            print(f"  3. 住宿 (zhusu)")
            print(f"  4. 滴滴/交通 (didi) - 包含地铁、公交、出租车等")
            print(f"  5. 退改签 (tuigai)")
            print(f"  6. 其他/跳过 (skip)")

            while True:
                try:
                    choice = input(f"请输入选项 (1-6) [默认4-滴滴/交通]: ").strip()
                    if choice == '' or choice == '4':
                        result = 'didi'
                        break
                    elif choice == '1':
                        result = 'jipiao'
                        break
                    elif choice == '2':
                        result = 'huoche'
                        break
                    elif choice == '3':
                        result = 'zhusu'
                        break
                    elif choice == '5':
                        result = 'tuigai'
                        break
                    elif choice == '6':
                        result = 'unknown'
                        break
                    else:
                        print("  无效选项，请重新输入")
                except (KeyboardInterrupt, EOFError):
                    print("\n  用户取消，跳过此发票")
                    return None

            # 记住用户选择（用于同一会话中的重复文件）
            USER_CONFIRMED_TYPES[filename] = result
            print(f"  用户选择: {result}")
            return result
        else:
            # 非交互式模式，默认归入滴滴类别
            print(f"  ⚠️ 无法识别发票类型，默认归入滴滴/交通类别: {filename}")
            return 'didi_other'

    return max_type


def is_toll_invoice_text(text):
    """判断文本是否为高速通行费发票或通行费电子行程单。"""
    if any(k in text for k in ['通行费电子行程单', '代收通行费', '收费公路通行费', '高速公路通行费', 'ETC']):
        return True
    if '通行费' in text and ('发票号码' in text or '价税合计' in text):
        return True
    station_keywords = sum(1 for k in ['入口站', '出口站', '入口时间', '出口时间', '交易金额'] if k in text)
    if station_keywords >= 3:
        return True
    if '入口：' in text and '出口：' in text and '通行费' in text:
        return True
    if '广东联合电子服务股份有限公司' in text and '行程信息' in text:
        return True
    return False


def classify_didi_subtype(text):
    """区分滴滴电子发票和行程单"""
    if '电子发票' in text and '发票号码' in text:
        return 'didi_einvoice'
    elif '行程单' in text or 'TRIP TABLE' in text.upper():
        return 'didi_trip'
    return 'didi_other'


def classify_toll_subtype(text):
    """区分高速费电子发票和通行费行程单"""
    if '通行费电子行程单' in text or ('入口站' in text and '出口站' in text and '交易金额' in text):
        return 'toll_trip'
    if '电子发票' in text and '发票号码' in text:
        return 'toll_einvoice'
    return 'toll_einvoice'


def build_invoice_rename(ftype, count):
    """根据发票类型和序号生成规范文件名"""
    if ftype == 'tuigai':
        return f"tuigai{count}.pdf"
    if ftype == 'huoche':
        return f"huoche{count}.pdf"
    if ftype == 'jipiao':
        return f"jipiao{count}.pdf"
    if ftype == 'zhusu':
        return f"zhusu{count}.pdf"
    if ftype == 'didi_einvoice':
        letter = chr(ord('A') + count - 1)
        return f"滴滴电子发票{letter}.pdf"
    if ftype == 'didi_trip':
        letter = chr(ord('A') + count - 1)
        return f"滴滴出行行程报销单{letter}.pdf"
    if ftype == 'didi_other':
        letter = chr(ord('A') + count - 1)
        return f"交通费发票{letter}.pdf"
    if ftype == 'toll_einvoice':
        letter = chr(ord('A') + count - 1)
        return f"高速费发票{letter}.pdf"
    if ftype == 'toll_trip':
        letter = chr(ord('A') + count - 1)
        return f"高速费行程单{letter}.pdf"
    if ftype == 'unknown':
        return f"未分类_{count}.pdf"
    return f"其他_{count}.pdf"


# ============ 城市提取函数 ============

def extract_cities_from_pdf(pdf_path, inv_type):
    """
    从机票或火车票PDF中提取出发站和到达站城市名称（排除北京）
    优先使用pdfplumber提取文本，失败后再使用OCR
    返回: (cities_list, raw_text)
    """
    text = ""

    # 优先使用pdfplumber提取文本
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
    except Exception as e:
        pass

    # 如果文本层为空或太短，使用OCR
    if len(text.strip()) < 50:
        images = pdf2image.convert_from_path(pdf_path, dpi=OCR_DPI_HIGH)
        text = pytesseract.image_to_string(images[0], lang='chi_sim')

    cities = set()

    def normalize_city_name(city):
        """标准化城市名：去掉站字和方向后缀"""
        if not city:
            return city
        # 去掉站字后缀
        city = re.sub(r'站$', '', city.strip())
        # 去掉东南西北方向后缀（如成都东->成都，北京南->北京）
        city = re.sub(r'[东南西北]$', '', city.strip())
        return city

    def is_valid_city_token(city, extra_blocklist=None):
        """过滤明显不是城市的 token。"""
        if not city:
            return False
        city = city.strip()
        if not city or len(city) > 10:
            return False
        if re.fullmatch(r'[0-9.]+', city):
            return False
        if re.search(r'\d', city):
            return False
        if city.upper() == 'CNY':
            return False
        blocked = {
            '城市', '地点', '目的', '出发', '到达', '自', '至',
            '国际', '国内', '大型', '四川省', '税务局', '发票号', '开票日'
        }
        if extra_blocklist:
            blocked.update(extra_blocklist)
        return city not in blocked

    if inv_type == 'jipiao':
        # 机票：查找出发和到达城市
        # 常见格式："自: 绵阳 南郊" "至: 北京 大兴" 或 "自: PKX 北京" "至:CTU 成都"

        for line in text.split('\n'):
            # 匹配 "自: 绵阳 南郊" 或 "自: PKX 北京" 这种格式
            if '自' in line and ':' in line:
                # 尝试匹配 "自: 词1 词2"，优先取中文词2（如果有的话）
                match = re.search(r'自\s*[:：]\s*(\S+)\s+(\S+)', line)
                if match:
                    word1, word2 = match.group(1).strip(), match.group(2).strip()
                    # 如果第一个词全是英文字母（机场代码），使用第二个词
                    if word1.isalpha() and word1.isascii():
                        city = normalize_city_name(word2)
                    else:
                        city = normalize_city_name(word1)
                    if is_valid_city_token(city):
                        cities.add(city)
                else:
                    # 只有一个词的情况
                    match = re.search(r'自\s*[:：]\s*(\S+)', line)
                    if match:
                        city = normalize_city_name(match.group(1).strip())
                        if is_valid_city_token(city):
                            cities.add(city)

            # 匹配 "至: 北京 大兴" 或 "至:CTU 成都" 这种格式
            if '至' in line and ':' in line:
                # 尝试匹配 "至: 词1 词2"，优先取中文词2（如果有的话）
                match = re.search(r'至\s*[:：]\s*(\S+)\s+(\S+)', line)
                if match:
                    word1, word2 = match.group(1).strip(), match.group(2).strip()
                    # 如果第一个词全是英文字母（机场代码），使用第二个词
                    if word1.isalpha() and word1.isascii():
                        city = normalize_city_name(word2)
                    else:
                        city = normalize_city_name(word1)
                    if is_valid_city_token(city):
                        cities.add(city)
                else:
                    # 只有一个词的情况
                    match = re.search(r'至\s*[:：]\s*(\S+)', line)
                    if match:
                        city = normalize_city_name(match.group(1).strip())
                        if is_valid_city_token(city):
                            cities.add(city)

        # 其他常见格式
        patterns = [
            r'出发[城市站]\s*(\S+?)\s*到达[城市站]\s*(\S+?)',
            r'从\s*(\S+?)\s*到\s*(\S+)',
            r'出发地[：:]\s*(\S+?)\s*目的地[：:]\s*(\S+)',
            r'(\S+?)[\s]*飞[\s]*(\S+)',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    for city in match:
                        city = city.strip()
                        if is_valid_city_token(city):
                            cities.add(city)
                else:
                    city = match.strip()
                    if is_valid_city_token(city):
                        cities.add(city)

        # 额外处理：查找常见的机场城市名称（基于上下文）
        # 查找 "XX机场" 格式的城市
        airport_pattern = r'(\S+?)(?:国际|国内)?机场'
        airport_matches = re.findall(airport_pattern, text)
        for city in airport_matches:
            city = normalize_city_name(city.strip())
            if is_valid_city_token(city, {'国际', '国内', '大型'}):
                cities.add(city)

    elif inv_type == 'huoche':
        # 火车票：查找出发站和到达站
        # 常见格式："出发站 北京南站 到达站 成都东站" 或 "北京南 - 成都东"
        patterns = [
            r'出发[站]\s*(\S+?)\s*到达[站]\s*(\S+?)',
            r'(\S+?站)\s*[-—–~～]\s*(\S+?站)',
            r'发[站]\s*[:：]?\s*(\S+?)\s*到[站]\s*[:：]?\s*(\S+)',
            # 移除过于宽泛的pattern：r'(\S+?)\s*[-—–~～]\s*(\S+?)'
            r'站名[：:]\s*(\S+?)\s*',
            r'(\S+?站)\s*[→到]\s*(\S+?站)',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    for station in match:
                        city = normalize_city_name(station.strip())
                        if is_valid_city_token(city):
                            cities.add(city)
                else:
                    city = normalize_city_name(match.strip())
                    if is_valid_city_token(city):
                        cities.add(city)

        # 查找 "站名" 后面的车站名称
        station_name_pattern = r'站名[：:]\s*(\S+?)(?:\s|$)'
        station_matches = re.findall(station_name_pattern, text)
        for station in station_matches:
            city = normalize_city_name(station.strip())
            if is_valid_city_token(city):
                cities.add(city)

        # 额外处理：处理两行格式的火车站信息（如"绵阳\n站"）
        # 匹配中文城市名，然后检查后续行是否包含"站"
        # 更严格的条件：城市名必须在行首，且附近有"站"
        lines = text.split('\n')
        for i, line in enumerate(lines):
            # 查找2-3个汉字的城市名（必须在行首或独立存在）
            city_matches = re.finditer(r'^([\u4e00-\u9fa5]{2,3})|(?<!\w)([\u4e00-\u9fa5]{2,3})(?!\w)', line)
            for city_match in city_matches:
                city_name = city_match.group(1) if city_match.group(1) else city_match.group(2)
                # 检查同一行或下一行是否包含"站"
                same_line_has_zhan = '站' in line
                next_line_has_zhan = i + 1 < len(lines) and '站' in lines[i + 1]
                # 只有附近有"站"才认为是车站城市
                if (same_line_has_zhan or next_line_has_zhan):
                    city = normalize_city_name(city_name)
                    if is_valid_city_token(city):
                        cities.add(city)

    # 过滤掉"北京"及其变体
    filtered_cities = set()
    for city in cities:
        # 排除北京及其常见变体
        if city not in ['北京', '北京市', '北京南', '北京北', '北京东', '北京西',
                        '北京南站', '北京北站', '北京东站', '北京西站', '首都']:
            filtered_cities.add(city)

    return list(filtered_cities), text


def get_unit_for_city(city, external_map=None):
    """
    根据城市名称返回对应的单位
    参数:
        city: 城市名称
        external_map: 外部传入的映射（如从配置文件或命令行参数加载）
    返回: (unit_name, needs_user_input)
    """
    # 使用外部传入的映射（由调用方从配置文件或命令行参数加载）
    city_unit_map = external_map or {}

    # 精确匹配
    if city in city_unit_map:
        return city_unit_map[city], False

    # 部分匹配（例如"成都市"匹配"成都"）
    for city_key, unit in city_unit_map.items():
        if city_key in city or city in city_key:
            return unit, False

    # 未知城市，需要用户输入
    return None, True


# ============ 数据提取函数 ============

def extract_amount_from_pdf(pdf_path, inv_type=None):
    """从PDF中提取价税合计金额 - 优先使用文本层，扫描件使用OCR"""
    text = ""

    # 先尝试提取PDF文本层
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
    except Exception as e:
        pass

    # 如果文本层为空或太短，使用OCR
    if len(text.strip()) < 50:
        images = pdf2image.convert_from_path(pdf_path, dpi=OCR_DPI_HIGH)
        text = pytesseract.image_to_string(images[0], lang='chi_sim')

    # 对于机票，优先取最大CNY金额（合计金额通常是最大的）
    if inv_type == 'jipiao':
        cny_pattern = r'CNY\s+([\d,]+\.\d{2})'
        cny_matches = re.findall(cny_pattern, text)
        if cny_matches:
            valid_numbers = [float(m.replace(',', '')) for m in cny_matches]
            if valid_numbers:
                return max(valid_numbers), '最大CNY金额(机票合计)'

    # 匹配金额模式（按优先级）
    patterns = [
        (r'价税合计.*?[（(]\s*小写\s*[）)].*?[¥￥]?\s*([\d,]+\.\d{2})', '价税合计(小写)'),
        (r'价税合计.*?[¥￥]\s*([\d,]+\.\d{2})', '价税合计'),
        (r'价税合计.*?([\d,]+\.\d{2})', '价税合计(纯数字)'),
        (r'退票费.*?[¥￥]\s*([\d,]+\.\d{2})', '退票费(后)'),
        (r'[¥￥]\s*([\d,]+\.\d{2})\s*退票费', '退票费(前)'),
        (r'改签费.*?[¥￥]\s*([\d,]+\.\d{2})', '改签费(后)'),
        (r'[¥￥]\s*([\d,]+\.\d{2})\s*改签费', '改签费(前)'),
        (r'合\s*计.*?CNY\s*([\d,]+\.\d{2})', '合计CNY'),
        (r'合\s*计.*?[¥￥]\s*([\d,]+\.\d{2})', '合计'),
    ]

    for pattern, desc in patterns:
        matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
        if matches:
            amount_str = matches[-1].replace(',', '').replace('¥', '').replace('￥', '')
            try:
                return float(amount_str), desc
            except:
                continue

    # 通用CNY匹配
    cny_pattern = r'[¥￥]\s*([\d,]+\.\d{2})'
    cny_matches = re.findall(cny_pattern, text)

    if cny_matches:
        valid_numbers = []
        for m in cny_matches:
            val = float(m.replace(',', ''))
            if 10 < val < 100000:
                valid_numbers.append(val)
        if valid_numbers:
            return max(valid_numbers), '最大CNY金额'

    # 最后的备用：找所有金额数字
    numbers = re.findall(r'[\d,]+\.\d{2}', text)
    if numbers:
        valid_numbers = []
        for n in numbers:
            val = float(n.replace(',', ''))
            if 10 < val < 100000:
                valid_numbers.append(val)
        if valid_numbers:
            return max(valid_numbers), '最大金额数字'

    return None, '未找到'


def extract_dates_with_pdfplumber(pdf_path, exclude_tiankai=True):
    """
    使用pdfplumber从PDF文本层提取日期
    用于OCR识别失败时的自动复核
    关键: 默认排除填开日期，只保留航班日期
    返回: (dates_list, text)
    """
    dates = []
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        # 提取日期模式
        date_patterns = [
            r'(\d{4}年\d{1,2}月\d{1,2}日)',
            r'(\d{4}-\d{2}-\d{2})',
            r'(\d{4}/\d{2}/\d{2})',
        ]

        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                try:
                    if '年' in match:
                        d = datetime.strptime(match, '%Y年%m月%d日')
                    elif '-' in match:
                        d = datetime.strptime(match, '%Y-%m-%d')
                    elif '/' in match:
                        d = datetime.strptime(match, '%Y/%m/%d')
                    dates.append(d)
                except:
                    continue

        unique_dates = sorted(list(set(dates)))

        # 排除填开日期
        if exclude_tiankai and unique_dates:
            tiankai_pattern = r'填开日期[:：]\s*(\d{4}年\d{1,2}月\d{1,2}日)'
            tiankai_matches = set(re.findall(tiankai_pattern, text))

            filtered_dates = []
            for d in unique_dates:
                d_str = d.strftime('%Y年%m月%d日')
                if d_str not in tiankai_matches:
                    filtered_dates.append(d)

            # 如果过滤后还有日期，使用过滤后的；否则返回原始日期（避免误伤）
            if filtered_dates:
                unique_dates = filtered_dates

        return unique_dates, text
    except Exception as e:
        return [], text


def extract_flight_dates_with_validation(pdf_path):
    """
    从机票PDF中提取航班日期，带多重校验和自动复核
    优先使用pdfplumber提取文本，失败后再使用OCR
    关键区分：航班日期 vs 填开日期
    返回: (dates_list, warnings_list, raw_text)
    """
    warnings = []

    # 优先使用pdfplumber提取文本
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
    except Exception as e:
        pass

    # 如果文本层为空或太短，使用OCR
    use_ocr = False
    if len(text.strip()) < 50:
        print(f"    ⚠️ PDF文本层不足，使用OCR识别...")
        images = pdf2image.convert_from_path(pdf_path, dpi=OCR_DPI_HIGH)
        text = pytesseract.image_to_string(images[0], lang='chi_sim')
        use_ocr = True
    else:
        print(f"    ✓ 使用PDF文本层识别")

    # 提取所有可能的日期（仅在OCR模式下处理"昌"->"日"错误）
    date_pattern = r'(\d{4}年\d{1,2}月\d{1,2}[日昌]?)'
    all_date_matches = re.findall(date_pattern, text)

    all_dates = []
    for match in all_date_matches:
        try:
            match_clean = match.replace('昌', '日').replace('日', '')
            d = datetime.strptime(match_clean, '%Y年%m月%d')
            all_dates.append(d)
        except:
            continue

    # ===== 自动复核：如果未识别到日期，尝试其他方法 =====
    if not all_dates:
        if use_ocr:
            print(f"    ⚠️ OCR未识别到日期，使用pdfplumber复核...")
            plumber_dates, plumber_text = extract_dates_with_pdfplumber(pdf_path)
            if plumber_dates:
                all_dates = plumber_dates
                text = plumber_text  # 使用pdfplumber的文本用于后续处理
                print(f"    ✓ pdfplumber复核成功，识别到 {len(plumber_dates)} 个日期")
        else:
            print(f"    ⚠️ PDF文本层未识别到日期，尝试OCR复核...")
            images = pdf2image.convert_from_path(pdf_path, dpi=OCR_DPI_HIGH)
            ocr_text = pytesseract.image_to_string(images[0], lang='chi_sim')
            ocr_date_matches = re.findall(date_pattern, ocr_text)
            for match in ocr_date_matches:
                try:
                    match_clean = match.replace('昌', '日').replace('日', '')
                    d = datetime.strptime(match_clean, '%Y年%m月%d')
                    all_dates.append(d)
                except:
                    continue
            if all_dates:
                text = ocr_text  # 使用OCR文本用于后续处理
                print(f"    ✓ OCR复核成功，识别到 {len(all_dates)} 个日期")

    if not all_dates:
        warnings.append(f"未识别到任何日期")
        return [], warnings, text

    # 去重并排序
    unique_dates = sorted(list(set(all_dates)))

    # ===== 关键：区分航班日期和填开日期 =====
    # 航班日期：在"日期"列下方，格式通常靠近"时间"列
    # 填开日期：通常在页面底部，标注为"填开日期"

    flight_dates = []

    # 方法1: 查找"日期"标题列下方的日期
    # 在航空行程单中，航班日期在行中，后面通常跟着时间（如 12:45）
    flight_date_pattern = r'(\d{4}年\d{1,2}月\d{1,2}[日昌]?)\s*\d{1,2}:\d{2}'
    flight_matches = re.findall(flight_date_pattern, text)
    for match in flight_matches:
        try:
            match_clean = match.replace('昌', '日').replace('日', '')
            d = datetime.strptime(match_clean, '%Y年%m月%d')
            flight_dates.append(d)
        except:
            pass

    # 方法2: 如果没有找到，分析文本结构
    # 航班日期通常在"日期"标题下方的第一行数据区域
    if not flight_dates:
        # 查找"日期"和"时间"之间的日期
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if '日期' in line and '时间' in line:
                # 在下一行查找日期
                if i + 1 < len(lines):
                    next_line = lines[i + 1]
                    date_match = re.search(r'(\d{4}年\d{1,2}月\d{1,2}[日昌]?)', next_line)
                    if date_match:
                        try:
                            match_clean = date_match.group(1).replace('昌', '日').replace('日', '')
                            d = datetime.strptime(match_clean, '%Y年%m月%d')
                            flight_dates.append(d)
                        except:
                            pass

    # 方法3: 如果前两种方法都没找到，使用所有日期（后续会统一过滤填开日期）
    if not flight_dates and unique_dates:
        flight_dates = unique_dates.copy()

    # ===== 关键步骤：统一过滤填开日期 =====
    # 查找"填开日期"关键字，排除它后面的日期
    tiankai_pattern = r'填开日期.*?(\d{4}年\d{1,2}月\d{1,2}[日昌]?)'
    tiankai_matches = set(re.findall(tiankai_pattern, text))

    if tiankai_matches and flight_dates:
        filtered_flight_dates = []
        for d in flight_dates:
            d_str = d.strftime('%Y年%m月%d日')
            d_str2 = d.strftime('%Y年%m月%d昌')
            if d_str not in tiankai_matches and d_str2 not in tiankai_matches:
                filtered_flight_dates.append(d)
            else:
                print(f"    📝 排除填开日期: {d.strftime('%Y-%m-%d')}")
        flight_dates = filtered_flight_dates

    # 如果过滤后没有日期了，使用原始日期（避免误过滤导致无日期）
    if not flight_dates and unique_dates:
        print(f"    ⚠️ 过滤后无日期，使用原始最早日期: {unique_dates[0].strftime('%Y-%m-%d')}")
        flight_dates = [unique_dates[0]]

    # ===== 自动复核：如果OCR只识别到1个日期但pdfplumber识别到多个，进行交叉验证 =====
    if len(flight_dates) == 1 and len(unique_dates) >= 2:
        plumber_dates, _ = extract_dates_with_pdfplumber(pdf_path)
        if len(plumber_dates) >= 2:
            print(f"    ⚠️ OCR只识别到1个日期但pdfplumber识别到{len(plumber_dates)}个，进行交叉验证...")
            # 使用pdfplumber的日期补充
            for pd in plumber_dates:
                if pd not in flight_dates:
                    flight_dates.append(pd)
            flight_dates = sorted(flight_dates)
            print(f"    ✓ 交叉验证后共有 {len(flight_dates)} 个航班日期")

    # ===== 校验逻辑 =====

    # 校验1: 检查是否有1月/2月日期（可能是11月/12月误识别）
    # 使用pdfplumber进行复核
    winter_date_warnings = []
    for d in flight_dates:
        if d.month in [1, 2]:
            winter_date_warnings.append((d, f"发现{d.month}月日期({d.strftime('%Y-%m-%d')})，可能是{d.month+10}月误识别，请人工复核"))

    # 对冬季日期进行pdfplumber复核
    if winter_date_warnings:
        plumber_dates, _ = extract_dates_with_pdfplumber(pdf_path)
        for d, warning in winter_date_warnings:
            # 检查pdfplumber是否识别到不同的月份
            for pd in plumber_dates:
                if pd.day == d.day and pd.year == d.year and pd.month != d.month:
                    if pd.month == d.month + 10:  # 11月 vs 1月, 12月 vs 2月
                        print(f"    ⚠️ 发现可能的OCR误识别: OCR识别为{d.month}月，pdfplumber识别为{pd.month}月")
                        print(f"    ✓ 自动修正日期: {d.strftime('%Y-%m-%d')} -> {pd.strftime('%Y-%m-%d')}")
                        flight_dates.remove(d)
                        flight_dates.append(pd)
                        break
            else:
                # 没有冲突，保留原警告
                warnings.append(warning)
        flight_dates = sorted(flight_dates)

    # 校验2: 检查日期是否合理
    now = datetime.now()
    for d in flight_dates:
        if d > now.replace(year=now.year + 1):
            warnings.append(f"日期{d.strftime('%Y-%m-%d')}超过当前时间+1年，可能识别错误")
        if d < now.replace(year=now.year - 2):
            warnings.append(f"日期{d.strftime('%Y-%m-%d')}早于当前时间-2年，可能识别错误")

    # 校验3: 如果有多个机票，检查日期跨度
    if len(flight_dates) >= 2:
        sorted_dates = sorted(flight_dates)
        span = (sorted_dates[-1] - sorted_dates[0]).days
        if span > 365:
            warnings.append(f"机票日期跨度{span}天超过1年，请检查")

    return flight_dates, warnings, text


# ============ 校验函数 ============

def validate_amount(inv_type, amount, filename):
    """校验金额合理性"""
    if amount is None:
        return False, "未提取到金额"

    min_val, max_val = AMOUNT_RANGES.get(inv_type, (0, 999999))

    if amount < min_val:
        return False, f"金额¥{amount}低于合理范围(¥{min_val}-{max_val})"
    if amount > max_val:
        return False, f"金额¥{amount}高于合理范围(¥{min_val}-{max_val})"

    return True, "金额合理"


def validate_flight_dates(dates_list, all_warnings):
    """校验机票日期集合的合理性"""
    warnings = list(all_warnings)

    if not dates_list:
        warnings.append("没有识别到任何航班日期")
        return False, warnings

    if len(dates_list) == 1:
        return True, warnings

    # 多日期排序检查
    sorted_dates = sorted(dates_list)

    # 检查时间跨度是否合理
    date_span = (sorted_dates[-1] - sorted_dates[0]).days
    if date_span > 365:
        warnings.append(f"航班日期跨度{date_span}天超过1年，请检查是否有误识别")

    return True, warnings


# ============ 主处理函数 ============

def process_invoices(input_dir, output_excel=None, sheet_name='sheet1', rename_files=True, force_write=False):
    """
    主处理函数

    参数:
        input_dir: 发票文件夹路径
        output_excel: 输出Excel文件路径（可选）
        sheet_name: 工作表名称
        rename_files: 是否重命名文件

    返回:
        dict: 包含处理结果和校验信息
    """

    result = {
        'success': True,
        'files': {},
        'summary': {},
        'warnings': [],
        'errors': [],
        'data': {}
    }

    # ========== 步骤0: 自动转换OFD文件 ==========
    ofd_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.ofd')]
    if ofd_files:
        print(f"[步骤0] 发现 {len(ofd_files)} 个OFD文件，开始转换...")

        # 使用 ofd2pdf 命令转换（需在PATH中，如 ~/apps/ofd2pdf）
        ofd2pdf_bin = shutil.which('ofd2pdf')
        if not ofd2pdf_bin:
            print("  ✗ 错误：未找到 ofd2pdf 命令")
            print("  原因：ofd2pdf 未安装或不在 PATH 中")
            print("  解决：请安装 ofd2pdf 并确保其可在终端直接执行，然后重新运行")
            sys.exit(1)
        else:
            original_cwd = os.getcwd()
            try:
                os.chdir(input_dir)
                ofd_args = ' '.join(f'"{f}"' for f in ofd_files)
                conv_result = subprocess.run(
                    f'ofd2pdf {ofd_args}',
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                os.chdir(original_cwd)

                if conv_result.returncode == 0:
                    for ofd_file in ofd_files:
                        pdf_name = os.path.splitext(ofd_file)[0] + '.pdf'
                        pdf_path = os.path.join(input_dir, pdf_name)
                        if os.path.exists(pdf_path):
                            print(f"  ✓ {ofd_file} -> {pdf_name}")
                        else:
                            print(f"  ✗ 错误：{ofd_file} 转换后未生成对应PDF")
                            print(f"  原因：ofd2pdf 执行成功但输出文件缺失，文件可能已损坏或格式不支持")
                            print(f"  解决：请手动转换该文件后重新运行")
                            sys.exit(1)
                else:
                    print(f"  ✗ 错误：OFD文件转换失败")
                    print(f"  原因：{conv_result.stderr.strip() or '未知错误（ofd2pdf 返回非零退出码）'}")
                    print(f"  解决：请检查上述错误信息，或手动将OFD文件转换为PDF后重新运行")
                    sys.exit(1)
            except subprocess.TimeoutExpired:
                os.chdir(original_cwd)
                print(f"  ✗ 错误：OFD转换超时（超过60秒）")
                print(f"  原因：文件过大或 ofd2pdf 程序无响应")
                print(f"  解决：请手动转换OFD文件后重新运行")
                sys.exit(1)
            except Exception as e:
                os.chdir(original_cwd)
                print(f"  ✗ 错误：OFD转换异常：{e}")
                print(f"  解决：请检查上述错误信息后重新运行")
                sys.exit(1)
        print()

    # 收集所有PDF文件
    pdf_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.pdf')]

    if not pdf_files:
        result['errors'].append(f"目录 {input_dir} 中没有PDF文件")
        return result

    print(f"发现 {len(pdf_files)} 个PDF文件，开始处理...")

    # 步骤1: 识别类型
    print("\n[步骤1] 识别发票类型...")
    file_types = {}
    file_texts = {}

    for filename in sorted(pdf_files):
        filepath = os.path.join(input_dir, filename)
        text = extract_text_from_pdf(filepath)
        file_texts[filename] = text

        inv_type = identify_invoice_type(text, filename, interactive=not force_write)

        # 用户取消，跳过此发票
        if inv_type is None:
            print(f"  {filename} -> 已跳过（用户取消）")
            continue

        # 用户选择跳过（其他/跳过）
        if inv_type == 'unknown':
            print(f"  {filename} -> unknown（用户选择跳过）")
            file_types[filename] = inv_type
            continue

        # 进一步细分滴滴类型
        if inv_type == 'didi':
            inv_type = classify_didi_subtype(text)
        elif inv_type == 'toll':
            inv_type = classify_toll_subtype(text)

        file_types[filename] = inv_type
        print(f"  {filename} -> {inv_type}")

    # 步骤2: 重命名文件（可选）
    if rename_files:
        print("\n[步骤2] 重命名文件...")
        counters = defaultdict(int)

        rename_map = {}
        for old_name, ftype in file_types.items():
            counters[ftype] += 1
            new_name = build_invoice_rename(ftype, counters[ftype])

            rename_map[old_name] = new_name

        # 执行重命名（使用临时文件名避免冲突）
        # 第一步：所有文件先重命名为临时名称
        temp_map = {}
        for i, (old_name, new_name) in enumerate(rename_map.items()):
            old_path = os.path.join(input_dir, old_name)
            if os.path.exists(old_path) and old_name != new_name:
                temp_name = f"__temp_{i}_{old_name}"
                temp_path = os.path.join(input_dir, temp_name)
                os.rename(old_path, temp_path)
                temp_map[old_name] = temp_name

        # 第二步：从临时名称重命名为最终名称
        for old_name, new_name in rename_map.items():
            if old_name in temp_map:
                temp_path = os.path.join(input_dir, temp_map[old_name])
                new_path = os.path.join(input_dir, new_name)
                os.rename(temp_path, new_path)
                print(f"  {old_name} -> {new_name}")

        # 更新文件列表
        pdf_files = list(rename_map.values())
        file_types = {new: file_types[old] for old, new in rename_map.items()}

    # 步骤3: 提取数据并校验
    print("\n[步骤3] 提取数据并校验...")

    all_flight_dates = []
    all_warnings = []
    all_cities = set()  # 收集所有城市（去重）

    data = {
        'huoche': {'count': 0, 'amounts': []},
        'jipiao': {'count': 0, 'amounts': [], 'dates': []},
        'zhusu': {'count': 0, 'amounts': []},
        'didi_einvoice': {'count': 0, 'amounts': []},
        'didi_trip': {'count': 0, 'amounts': []},  # 行程报销单，不计入金额（与电子发票重复）
        'didi_other': {'count': 0, 'amounts': []},  # 其他交通费发票（地铁、公交等）
        'toll_einvoice': {'count': 0, 'amounts': []},
        'toll_trip': {'count': 0, 'amounts': []},  # 通行费行程单，不计入金额（与发票重复）
        'tuigai': {'count': 0, 'amounts': []},
    }

    # 文件明细数据（用于阶段3匹配交易单号）
    file_details = []

    for filename in pdf_files:
        filepath = os.path.join(input_dir, filename)
        ftype = file_types.get(filename, 'unknown')

        # 提取金额
        amount, source = extract_amount_from_pdf(filepath, ftype)

        # 记录文件明细（用于阶段3）
        file_info = {
            'filename': filename,
            'type': ftype,
            'amount': amount,
        }

        # 对于滴滴行程单，记录所有行程金额
        if ftype == 'didi_trip':
            trip_amounts = extract_amounts_from_didi_trip(filepath)
            if trip_amounts:
                file_info['trip_amounts'] = trip_amounts

        # 金额校验（跳过 unknown 类型）
        if amount and ftype != 'unknown':
            validate_type = (
                ftype.replace('didi_einvoice', 'didi')
                .replace('didi_trip', 'didi')
                .replace('didi_other', 'didi')
                .replace('toll_einvoice', 'toll')
                .replace('toll_trip', 'toll')
            )
            valid, msg = validate_amount(validate_type, amount, filename)
            if not valid:
                all_warnings.append(f"{filename}: {msg}")

        # 特殊处理机票日期和城市
        if ftype == 'jipiao':
            dates, warnings, raw_text = extract_flight_dates_with_validation(filepath)
            all_warnings.extend([f"{filename}: {w}" for w in warnings])

            if dates:
                data['jipiao']['dates'].extend(dates)
                all_flight_dates.extend(dates)

            # 提取城市
            cities, city_raw_text = extract_cities_from_pdf(filepath, 'jipiao')
            if cities:
                all_cities.update(cities)
                print(f"  {filename}: 金额¥{amount}, 航班日期: {[d.strftime('%Y-%m-%d') for d in dates]}, 城市: {cities}")
            else:
                # 再次确认日期（交叉验证）
                if len(dates) >= 1:
                    print(f"  {filename}: 金额¥{amount}, 航班日期: {[d.strftime('%Y-%m-%d') for d in dates]}")
                else:
                    print(f"  {filename}: 金额¥{amount}, 未识别到航班日期")

            # 日期合理性交叉校验
            if len(all_flight_dates) >= 2:
                sorted_dates = sorted(all_flight_dates)
                # 检查是否有异常早的日期（可能是误识别）
                if (sorted_dates[-1] - sorted_dates[0]).days > 300:
                    all_warnings.append(f"{filename}: 与其他机票日期跨度超过300天，请检查是否有1月/11月误识别")

        # 特殊处理火车票城市
        elif ftype == 'huoche':
            # 提取城市
            cities, city_raw_text = extract_cities_from_pdf(filepath, 'huoche')
            if cities:
                all_cities.update(cities)
                print(f"  {filename}: 金额¥{amount}, 城市: {cities}")
            else:
                print(f"  {filename}: 金额¥{amount}")

        else:
            print(f"  {filename}: 金额¥{amount}")

        # 累加数据（跳过 unknown 类型）
        if ftype == 'unknown':
            continue

        base_type = ftype
        if 'didi' in ftype:
            if 'einvoice' in ftype:
                base_type = 'didi_einvoice'
            elif 'trip' in ftype:
                base_type = 'didi_trip'
            else:
                base_type = 'didi_other'  # 其他交通费（地铁、公交等）
        elif 'toll' in ftype:
            if 'trip' in ftype:
                base_type = 'toll_trip'
            else:
                base_type = 'toll_einvoice'

        if base_type in data:
            data[base_type]['count'] += 1
            if amount and 'amounts' in data[base_type] and base_type != 'toll_trip':
                data[base_type]['amounts'].append(amount)

        # 添加到文件明细列表
        file_details.append(file_info)

    # 汇总统计
    print("\n[步骤4] 数据汇总...")

    # 计算总滴滴文件数（电子发票+行程单）
    total_didi = data['didi_einvoice']['count'] + data['didi_trip']['count']
    total_toll = data['toll_einvoice']['count'] + data['toll_trip']['count']

    # 保存日期对象供后续使用
    earliest_date_obj = min(data['jipiao']['dates']) if data['jipiao']['dates'] else None
    latest_date_obj = max(data['jipiao']['dates']) if data['jipiao']['dates'] else None

    # 处理城市信息：转换为列表并排序，用中文顿号连接
    sorted_cities = sorted(list(all_cities))
    cities_str = '、'.join(sorted_cities) if sorted_cities else ''
    print(f"  提取到的城市（除北京）: {sorted_cities}")
    print(f"  城市字符串: {cities_str}")

    summary = {
        'huoche_count': data['huoche']['count'],
        'huoche_amount': sum(data['huoche']['amounts']),
        'jipiao_count': data['jipiao']['count'],
        'jipiao_amount': sum(data['jipiao']['amounts']),
        'jipiao_earliest_date': earliest_date_obj.strftime('%m/%d/%Y') if earliest_date_obj else None,
        'jipiao_latest_date': latest_date_obj.strftime('%m/%d/%Y') if latest_date_obj else None,
        'tuigai_count': data['tuigai']['count'],
        'tuigai_amount': sum(data['tuigai']['amounts']),
        'didi_count': total_didi,
        'didi_einvoice_amount': sum(data['didi_einvoice']['amounts']) + sum(data['didi_other']['amounts']),  # 只计算电子发票+其他交通费，排除行程报销单（与发票重复）
        'toll_count': total_toll,
        'toll_amount': sum(data['toll_einvoice']['amounts']),
        'zhusu_amount': sum(data['zhusu']['amounts']),
        'cities': sorted_cities,
        'cities_str': cities_str,
    }

    # 保存日期对象供Excel写入使用
    result['date_objects'] = {
        'earliest': earliest_date_obj,
        'latest': latest_date_obj
    }

    for key, value in summary.items():
        print(f"  {key}: {value}")

    result['summary'] = summary
    result['warnings'] = all_warnings
    result['data'] = data
    result['file_details'] = file_details

    # 步骤5: 填写Excel（如果指定）
    if output_excel and os.path.exists(output_excel):
        print(f"\n[步骤5] 填写Excel表格: {output_excel}")

        # 显示提取的数据摘要，方便用户确认
        print("\n  提取的数据摘要:")
        print(f"    城市: {summary.get('cities_str', '无')}")
        print(f"    火车票: {summary['huoche_count']}张, 金额¥{summary['huoche_amount']}")
        print(f"    机票: {summary['jipiao_count']}张, 金额¥{summary['jipiao_amount']}")
        if summary['jipiao_earliest_date']:
            print(f"    最早日期: {summary['jipiao_earliest_date']}")
        if summary['jipiao_latest_date']:
            print(f"    最晚日期: {summary['jipiao_latest_date']}")
        print(f"    滴滴/交通: {summary['didi_count']}张, 金额¥{summary['didi_einvoice_amount']}")
        print(f"    高速费: {summary['toll_count']}张, 金额¥{summary['toll_amount']}")
        print(f"    住宿: 金额¥{summary['zhusu_amount']}")

        should_write = True

        if all_warnings:
            print("\n" + "="*60)
            print("⚠️  数据校验警告 - 需要人工复核")
            print("="*60)

            # 分类显示警告
            date_warnings = [w for w in all_warnings if '日期' in w or '月' in w]
            amount_warnings = [w for w in all_warnings if '金额' in w]
            other_warnings = [w for w in all_warnings if w not in date_warnings and w not in amount_warnings]

            if date_warnings:
                print("\n【日期问题 - 请务必复核】")
                for w in date_warnings:
                    print(f"  ❌ {w}")
                print("  💡 提示: OCR常将'11月'误识为'1月'，'29日'误识为'9日'")

            if amount_warnings:
                print("\n【金额问题】")
                for w in amount_warnings:
                    print(f"  ⚠️ {w}")

            if other_warnings:
                print("\n【其他问题】")
                for w in other_warnings:
                    print(f"  ⚠️ {w}")

            print("\n" + "="*60)
            print("【提取的数据摘要】")
            print(f"  城市: {summary.get('cities_str', '无')}")
            print(f"  机票: {summary['jipiao_count']}张, 金额¥{summary['jipiao_amount']}")
            if summary['jipiao_earliest_date']:
                print(f"  最早日期: {summary['jipiao_earliest_date']}")
            if summary['jipiao_latest_date']:
                print(f"  最晚日期: {summary['jipiao_latest_date']}")
            print(f"  住宿: 金额¥{summary['zhusu_amount']}")
            print(f"  滴滴/交通: {summary['didi_count']}张, 金额¥{summary['didi_einvoice_amount']}")
            print(f"  高速费: {summary['toll_count']}张, 金额¥{summary['toll_amount']}")
            print("="*60)

            # 如果使用了 --force-write，直接写入不询问
            if force_write:
                print("\n(使用 --force-write 参数，跳过确认直接写入)")
                should_write = True
            else:
                # 交互式确认（如果在终端运行）
                if sys.stdin.isatty():
                    print("\n请选择操作:")
                    print("  1. 确认数据正确，继续写入Excel")
                    print("  2. 跳过写入，手动修改后重新运行")
                    print("  3. 标记问题后继续（在Excel中手动修正）")
                    try:
                        response = input("\n请输入选项 (1/2/3, 默认1): ").strip()
                        if response in ('2', 'no', '否', 'n'):
                            should_write = False
                            print("  ❌ 已取消写入。请检查发票后重新运行。")
                            result['pending_write'] = True
                        elif response == '3':
                            should_write = True
                            print("  ⚠️ 将继续写入，请在Excel中手动修正数据")
                            result['needs_manual_fix'] = True
                        else:
                            # 默认选项1或输入1
                            should_write = True
                            print("  ✓ 确认继续写入")
                    except (EOFError, KeyboardInterrupt):
                        # 非交互式环境，默认继续写入
                        print("  (非交互式环境，默认继续写入)")
                        should_write = True
                else:
                    # 非终端环境，默认继续写入但标记警告
                    print("\n(非交互式环境，默认继续写入)")
                    print("  如需跳过确认，请使用 --force-write 参数")
                    should_write = True

        if should_write:
            write_to_excel(output_excel, sheet_name, summary, result.get('date_objects'))
            result['pending_write'] = False
            print("\n  ✓ 数据已写入Excel")
        else:
            print("\n  ⚠️ 数据未写入Excel")

    return result


def amount_to_chinese(amount):
    """
    将金额数字转换为中文大写金额（经过测试的实现）

    Args:
        amount: 数字金额（支持整数、小数、字符串）

    Returns:
        str: 中文大写金额

    Examples:
        >>> amount_to_chinese(1234.56)
        '壹仟贰佰叁拾肆元伍角陆分'
        >>> amount_to_chinese(100000)
        '壹拾万元整'
    """
    import re

    # 数字与大写映射
    num_map = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    unit_map = ['', '拾', '佰', '仟']
    big_unit = ['', '万', '亿', '兆']

    # 转换为字符串并清理
    amount_str = str(amount).replace(',', '').replace(' ', '')

    # 验证格式
    if not re.match(r'^-?\d+\.?\d*$', amount_str):
        raise ValueError(f"无效金额格式: {amount}")

    # 处理负数
    negative = amount_str.startswith('-')
    if negative:
        amount_str = amount_str[1:]

    # 分离整数和小数部分
    if '.' in amount_str:
        integer_part, decimal_part = amount_str.split('.')
        decimal_part = decimal_part[:2]  # 最多两位小数
    else:
        integer_part, decimal_part = amount_str, ''

    # 验证金额范围
    if len(integer_part) > 16:
        raise ValueError("金额超出支持范围（最大支持到千兆）")

    result = []

    # 转换整数部分
    if integer_part == '0' or integer_part == '':
        if not decimal_part:
            return '零元整'
    else:
        # 去除前导零
        integer_part = integer_part.lstrip('0') or '0'

        # 按4位分组（从右向左）
        groups = []
        temp = integer_part
        while temp:
            groups.insert(0, temp[-4:].zfill(4))
            temp = temp[:-4]

        zero_flag = False

        for i, group in enumerate(groups):
            group_int = int(group)

            if group_int == 0:
                zero_flag = True
                continue

            # 处理这个4位组
            group_str = group.zfill(4)
            group_result = []
            has_nonzero = False

            for j, digit in enumerate(group_str):
                d = int(digit)
                if d == 0:
                    if has_nonzero and (not group_result or group_result[-1] != '零'):
                        group_result.append('零')
                else:
                    has_nonzero = True
                    group_result.append(num_map[d] + unit_map[3-j])

            # 去除末尾的零
            while group_result and group_result[-1] == '零':
                group_result.pop()

            # 添加大单位
            if group_result:
                if zero_flag:
                    result.append('零')
                result.extend(group_result)
                unit_idx = len(groups) - 1 - i
                if unit_idx > 0:
                    result.append(big_unit[unit_idx])
                zero_flag = False

        # 添加"元"
        if result and result[-1] == '零':
            result[-1] = '元'
        else:
            result.append('元')

    # 转换小数部分
    if decimal_part:
        jiao = int(decimal_part[0]) if len(decimal_part) >= 1 else 0
        fen = int(decimal_part[1]) if len(decimal_part) >= 2 else 0

        if jiao == 0 and fen == 0:
            result.append('整')
        else:
            if jiao > 0:
                result.append(num_map[jiao] + '角')
            elif integer_part != '0' and integer_part != '':
                result.append('零')

            if fen > 0:
                result.append(num_map[fen] + '分')
    else:
        result.append('整')

    # 清理连续的零
    result_str = ''.join(result)
    result_str = re.sub(r'零+', '零', result_str)
    result_str = result_str.replace('零元', '元')
    result_str = result_str.replace('零整', '整')

    # 处理负数
    if negative:
        result_str = '负' + result_str

    return result_str


def write_to_excel(excel_path, sheet_name, summary, date_objects=None):
    """将汇总数据写入Excel表格

    参数:
        date_objects: 包含datetime对象的字典 {'earliest': datetime, 'latest': datetime}
                     如果提供，则写入日期对象并设置中文显示格式
    """

    wb = load_workbook(excel_path)
    ws = wb[sheet_name]

    # 计算F10单元格的值：市内交通费包含滴滴/其他交通和高速费发票金额。
    local_transport_amount = summary['didi_einvoice_amount'] + summary.get('toll_amount', 0)
    local_transport_count = summary['didi_count'] + summary.get('toll_count', 0)
    if date_objects and date_objects.get('earliest') and date_objects.get('latest'):
        trip_days = (date_objects['latest'] - date_objects['earliest']).days + 1
        max_allowable = trip_days * 80
        f10_value = min(local_transport_amount, max_allowable)
        print(f"\n  F10计算: 实际金额¥{local_transport_amount}, 差旅天数{trip_days}天, 上限¥{max_allowable}, 取较小值¥{f10_value}")
    else:
        f10_value = local_transport_amount

    updates = {
        'E4': summary.get('cities_str', ''),
        'E6': summary['huoche_count'],
        'F6': summary['huoche_amount'],
        'E7': summary['jipiao_count'],
        'F7': summary['jipiao_amount'],
        'E8': summary['tuigai_count'],
        'F8': summary['tuigai_amount'],
        'E10': local_transport_count,
        'F10': f10_value,
        'O9': summary['zhusu_amount'],
    }

    print("\n  写入数据:")
    for cell, value in updates.items():
        old_value = ws[cell].value
        ws[cell] = value
        print(f"    {cell}: {old_value} -> {value}")

    # 处理日期单元格：存储为日期对象，显示为中文格式
    if date_objects:
        # 中文日期显示格式：2025年11月26日
        chinese_date_format = 'yyyy"年"m"月"d"日"'

        if 'earliest' in date_objects and date_objects['earliest']:
            ws['J4'] = date_objects['earliest']
            ws['J4'].number_format = chinese_date_format
            print(f"    J4: 设置为日期 {date_objects['earliest'].strftime('%Y-%m-%d')}, 显示格式: yyyy年m月d日")

        if 'latest' in date_objects and date_objects['latest']:
            ws['M4'] = date_objects['latest']
            ws['M4'].number_format = chinese_date_format
            print(f"    M4: 设置为日期 {date_objects['latest'].strftime('%Y-%m-%d')}, 显示格式: yyyy年m月d日")
    else:
        # 兼容旧方式：直接写入字符串
        if summary['jipiao_earliest_date']:
            ws['J4'] = summary['jipiao_earliest_date']
        if summary['jipiao_latest_date']:
            ws['M4'] = summary['jipiao_latest_date']

    wb.save(excel_path)
    print(f"\n  已保存到: {excel_path}")


def copy_dates_between_sheets(excel_path, source_sheet='sheet1', target_sheet='sheet2'):
    """
    将日期从sheet1复制到sheet2的指定位置
    - sheet1的J4 -> sheet2的D3 (最早日期)
    - sheet1的M4 -> sheet2的F3 (最晚日期)
    """
    print(f"\n[步骤5.5] 复制日期到 {target_sheet}...")

    try:
        wb = load_workbook(excel_path)

        if source_sheet not in wb.sheetnames:
            print(f"  ⚠️ 源工作表 '{source_sheet}' 不存在")
            return False

        if target_sheet not in wb.sheetnames:
            print(f"  ⚠️ 目标工作表 '{target_sheet}' 不存在，跳过复制")
            return False

        ws_source = wb[source_sheet]
        ws_target = wb[target_sheet]

        # 复制J4到D3 (最早日期)
        j4_value = ws_source['J4'].value
        if j4_value:
            ws_target['D3'] = j4_value
            # 如果源单元格有日期格式，也复制格式
            if ws_source['J4'].number_format:
                ws_target['D3'].number_format = ws_source['J4'].number_format
            print(f"  ✓ {source_sheet}.J4 ({j4_value}) -> {target_sheet}.D3")

        # 复制M4到F3 (最晚日期)
        m4_value = ws_source['M4'].value
        if m4_value:
            ws_target['F3'] = m4_value
            # 如果源单元格有日期格式，也复制格式
            if ws_source['M4'].number_format:
                ws_target['F3'].number_format = ws_source['M4'].number_format
            print(f"  ✓ {source_sheet}.M4 ({m4_value}) -> {target_sheet}.F3")

        wb.save(excel_path)
        print(f"  已保存到: {excel_path}")
        return True

    except Exception as e:
        print(f"  ❌ 复制日期失败: {e}")
        return False


def convert_excel_to_pdf(excel_path, output_pdf):
    """将Excel文件转换为PDF（使用LibreOffice）"""
    print(f"\n[步骤6] 转换Excel为PDF: {excel_path} -> {output_pdf}")

    try:
        import subprocess
        import tempfile
        import shutil

        # 检查LibreOffice是否安装
        try:
            subprocess.run(['which', 'libreoffice'], check=True, capture_output=True)
        except subprocess.CalledProcessError:
            print(f"  ❌ LibreOffice未安装，请先安装: apt-get install libreoffice-writer libreoffice-calc")
            return None

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()

        # 设置环境变量以支持中文字体
        env = os.environ.copy()
        env['LANG'] = 'zh_CN.UTF-8'
        env['LC_ALL'] = 'zh_CN.UTF-8'

        # 使用libreoffice命令行转换
        # 添加 --infilter 参数确保正确处理中文字符
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', temp_dir,
            excel_path
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)

        if result.returncode == 0:
            # 找到生成的PDF文件
            excel_name = os.path.splitext(os.path.basename(excel_path))[0]
            generated_pdf = os.path.join(temp_dir, f"{excel_name}.pdf")

            if os.path.exists(generated_pdf):
                shutil.move(generated_pdf, output_pdf)
                print(f"  ✓ Excel已转换为PDF: {output_pdf}")
                return output_pdf
            else:
                # 查找任何生成的PDF
                pdf_files = [f for f in os.listdir(temp_dir) if f.endswith('.pdf')]
                if pdf_files:
                    shutil.move(os.path.join(temp_dir, pdf_files[0]), output_pdf)
                    print(f"  ✓ Excel已转换为PDF: {output_pdf}")
                    return output_pdf

        print(f"  ⚠️ LibreOffice转换失败: {result.stderr}")
        return None

    except Exception as e:
        print(f"  ❌ Excel转PDF失败: {e}")
        return None


def fill_docx_document(docx_path, cities, earliest_date, latest_date, city_unit_map=None, has_flight=False, has_train=False):
    """
    填写Word文档表格内容
    - 将城市名称填到"地点:"后面
    - 将"到达单位:"后面填写对应单位
    - 将起止时间修改为最早日期至最晚日期
    - 设置字体为小四(12pt)
    - 根据发票类型在出行方式前打勾

    参数:
        docx_path: Word文档路径
        cities: 城市列表
        earliest_date: 最早日期 (datetime对象或字符串)
        latest_date: 最晚日期 (datetime对象或字符串)
        city_unit_map: 城市到单位的映射字典，如 {'北京': '总部', '上海': '分公司'}
        has_flight: 是否有飞机票
        has_train: 是否有火车票
    """
    print(f"\n[步骤6.5] 填写Word文档: {docx_path}")

    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document(docx_path)

        # 小四字体大小 = 12pt
        XIAOSI_SIZE = Pt(12)

        def set_cell_format(cell, size=XIAOSI_SIZE, align_center=True):
            """设置单元格中所有段落的字体大小和对齐方式"""
            for paragraph in cell.paragraphs:
                if align_center:
                    paragraph.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER = 1
                for run in paragraph.runs:
                    run.font.size = size

        def format_transportation(text, has_flight, has_train):
            """根据发票类型设置出行方式打勾"""
            # 构建新的出行方式字符串
            parts = []

            if has_flight:
                parts.append('☑飞机')
            else:
                parts.append('□飞机')

            if has_train:
                parts.append('☑高铁/动车')
            else:
                parts.append('□高铁/动车')

            # 其他选项保持原样（未选中）
            parts.extend(['□普通列车', '□长途汽车', '□轮船', '□其他'])

            return '  '.join(parts)

        # 准备日期字符串
        if isinstance(earliest_date, datetime):
            earliest_str = earliest_date.strftime('%Y年%m月%d日')
        else:
            earliest_str = str(earliest_date) if earliest_date else ''

        if isinstance(latest_date, datetime):
            latest_str = latest_date.strftime('%Y年%m月%d日')
        else:
            latest_str = str(latest_date) if latest_date else ''

        date_range_str = f"{earliest_str}至{latest_str}"

        # 准备城市字符串（中文顿号连接）
        cities_str = '、'.join(sorted(cities)) if cities else ''

        # 准备单位字符串
        units = set()
        unknown_cities = []

        # 城市单位映射（由调用方从配置文件或命令行参数传入）
        effective_map = city_unit_map or {}

        for city in sorted(cities):
            unit, needs_input = get_unit_for_city(city, effective_map)
            if unit:
                units.add(unit)
            else:
                unknown_cities.append(city)

        # 如果有未知城市，提示用户输入
        if unknown_cities:
            print(f"\n  ⚠️ 发现未知城市需要确认单位: {unknown_cities}")
            # 这里返回False，让调用者处理用户输入
            return False, unknown_cities, cities_str

        units_str = '、'.join(sorted(units)) if units else ''

        print(f"  填写信息:")
        print(f"    城市: {cities_str}")
        print(f"    单位: {units_str}")
        print(f"    起止时间: {date_range_str}")

        # 遍历文档中的所有表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cells = list(row.cells)
                for cell_idx, cell in enumerate(cells):
                    cell_text = cell.text.strip()

                    # 处理"到达单位\n及地点"标题行 - 在第一列，填写到第二列
                    if '到达单位' in cell_text and '地点' in cell_text and cell_idx == 0:
                        # 找到标题行，在第二列填写完整内容
                        if len(cells) > 1:
                            target_cell = cells[1]
                            original_text = target_cell.text
                            # 组合到达单位和地点
                            new_text = f"到达单位：{units_str}\n地点：{cities_str}"
                            target_cell.text = new_text
                            set_cell_format(target_cell)
                            print(f"    表{table_idx}行{row_idx}列1: '{original_text}' -> '{new_text}'")

                    # 查找"地点:"单元格（单独的情况）
                    if '地点：' in cell_text or '地点:' in cell_text:
                        # 如果不是刚处理的组合单元格
                        if '到达单位' not in cell_text:
                            original_text = cell.text
                            # 保留原有的格式，只替换内容
                            if '地点：' in cell_text:
                                cell.text = f"地点：{cities_str}"
                            else:
                                cell.text = f"地点:{cities_str}"
                            print(f"    表{table_idx}行{row_idx}列{cell_idx}: '{original_text}' -> '{cell.text}'")

                    # 查找"到达单位:"单元格（单独的情况）
                    if ('到达单位：' in cell_text or '到达单位:' in cell_text) and '地点' not in cell_text:
                        original_text = cell.text
                        # 检查是否需要添加到现有内容后面
                        remaining = cell_text.replace('到达单位：', '').replace('到达单位:', '').strip()
                        if not remaining:
                            # 为空，直接填写
                            if '到达单位：' in cell_text:
                                cell.text = f"到达单位：{units_str}"
                            else:
                                cell.text = f"到达单位:{units_str}"
                            print(f"    表{table_idx}行{row_idx}列{cell_idx}: '{original_text}' -> '{cell.text}'")

                    # 查找"起止时间"单元格或行 - 只填写该行的日期
                    if '起止时间' in cell_text:
                        # 在同行查找日期内容并替换（只处理"起止时间"行）
                        # 从第1列开始填写（跳过第0列的"起止时间"标题）
                        for next_cell_idx in range(1, len(row.cells)):
                            next_cell = row.cells[next_cell_idx]
                            # 检查是否包含日期格式（年...至/到）
                            if '年' in next_cell.text and ('至' in next_cell.text or '到' in next_cell.text):
                                original_text = next_cell.text
                                next_cell.text = date_range_str
                                set_cell_format(next_cell)
                                print(f"    表{table_idx}行{row_idx}列{next_cell_idx}: 起止时间 '{original_text}' -> '{next_cell.text}'")

                    # 查找"出行方式"单元格并设置打勾
                    if '飞机' in cell_text and '高铁' in cell_text:
                        if has_flight or has_train:  # 只有在有发票数据时才处理
                            original_text = cell.text
                            new_text = format_transportation(cell_text, has_flight, has_train)
                            cell.text = new_text
                            set_cell_format(cell)
                            print(f"    表{table_idx}行{row_idx}列{cell_idx}: 出行方式 '{original_text}' -> '{new_text}'")

        # 保存修改后的文档
        doc.save(docx_path)
        print(f"  ✓ Word文档已保存: {docx_path}")
        return True, [], cities_str

    except ImportError:
        print("  ❌ 未安装python-docx库，无法填写Word文档")
        return False, [], ''
    except Exception as e:
        print(f"  ❌ 填写Word文档失败: {e}")
        return False, [], ''


def convert_docx_to_pdf(docx_path, output_pdf):
    """将Word文档转换为PDF（使用LibreOffice）"""
    print(f"\n[步骤7] 转换Word为PDF: {docx_path} -> {output_pdf}")

    try:
        import subprocess
        import tempfile
        import shutil

        # 检查LibreOffice是否安装
        try:
            subprocess.run(['which', 'libreoffice'], check=True, capture_output=True)
        except subprocess.CalledProcessError:
            print(f"  ❌ LibreOffice未安装，请先安装: apt-get install libreoffice-writer")
            return None

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()

        # 设置环境变量以支持中文字体
        env = os.environ.copy()
        env['LANG'] = 'zh_CN.UTF-8'
        env['LC_ALL'] = 'zh_CN.UTF-8'

        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', temp_dir,
            docx_path
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)

        if result.returncode == 0:
            docx_name = os.path.splitext(os.path.basename(docx_path))[0]
            generated_pdf = os.path.join(temp_dir, f"{docx_name}.pdf")

            if os.path.exists(generated_pdf):
                shutil.move(generated_pdf, output_pdf)
                print(f"  ✓ Word已转换为PDF: {output_pdf}")
                return output_pdf
            else:
                pdf_files = [f for f in os.listdir(temp_dir) if f.endswith('.pdf')]
                if pdf_files:
                    shutil.move(os.path.join(temp_dir, pdf_files[0]), output_pdf)
                    print(f"  ✓ Word已转换为PDF: {output_pdf}")
                    return output_pdf

        print(f"  ⚠️ LibreOffice转换失败: {result.stderr}")
        return None

    except Exception as e:
        print(f"  ❌ Word转PDF失败: {e}")
        return None


def merge_pdfs(pdf_files, output_path):
    """合并多个PDF文件"""
    print(f"\n[步骤8] 合并PDF文件为: {output_path}")
    print(f"  共 {len(pdf_files)} 个PDF文件")

    try:
        from pypdf import PdfMerger

        merger = PdfMerger()

        for pdf_file in pdf_files:
            if os.path.exists(pdf_file):
                merger.append(pdf_file)
                print(f"  ✓ 添加: {os.path.basename(pdf_file)}")
            else:
                print(f"  ⚠️ 文件不存在: {pdf_file}")

        merger.write(output_path)
        merger.close()

        print(f"  ✓ 合并完成: {output_path}")
        return output_path

    except ImportError:
        # 备用：使用PyPDF2
        try:
            from PyPDF2 import PdfMerger

            merger = PdfMerger()

            for pdf_file in pdf_files:
                if os.path.exists(pdf_file):
                    merger.append(pdf_file)
                    print(f"  ✓ 添加: {os.path.basename(pdf_file)}")

            merger.write(output_path)
            merger.close()

            print(f"  ✓ 合并完成: {output_path}")
            return output_path

        except Exception as e:
            print(f"  ❌ 合并失败: {e}")
            return None

    except Exception as e:
        print(f"  ❌ 合并失败: {e}")
        return None


def collect_pdfs_in_order(invoice_dir, work_dir):
    """按指定顺序收集PDF文件"""

    # 定义顺序和对应的文件名模式
    order_patterns = [
        ('biaoge', ['biaoge.pdf']),
        ('shenpi', ['shenpi.pdf']),
        ('jipiao', ['jipiao*.pdf']),
        ('huoche', ['huoche*.pdf']),
        ('tuigai', ['tuigai*.pdf']),
        ('zhusu', ['zhusu*.pdf']),
        ('didi', ['滴滴*.pdf', '交通费*.pdf']),
        ('toll', ['高速费*.pdf']),
    ]

    collected_files = []

    print("\n  按以下顺序收集PDF文件:")
    for category, patterns in order_patterns:
        print(f"\n  [{category}]")
        category_files = []

        for pattern in patterns:
            # 在工作目录查找
            if category in ['biaoge', 'shenpi']:
                search_path = work_dir
            else:
                search_path = invoice_dir

            import glob
            matches = glob.glob(os.path.join(search_path, pattern))
            matches.sort()  # 确保按文件名排序

            for match in matches:
                if os.path.exists(match):
                    category_files.append(match)
                    print(f"    - {os.path.basename(match)}")

        collected_files.extend(category_files)

    return collected_files


# ============ 命令行入口 ============

def generate_data_report(result, output_path='data_report.json', config=None):
    """生成数据报告（JSON格式），供用户复核

    参数:
        result: 处理结果字典
        output_path: 输出文件路径
        config: 配置字典，用于检测未知城市
    """
    import json
    from datetime import datetime

    summary = result.get('summary', {})
    date_objects = result.get('date_objects', {})
    cities = summary.get('cities', [])

    # 检测未知城市（不在config.yaml中的城市）
    unknown_cities = []
    if config and 'city_units' in config:
        city_unit_map = config['city_units']
        for city in cities:
            # 检查城市是否在配置中
            is_known = False
            if city in city_unit_map:
                is_known = True
            else:
                # 部分匹配
                for known_city in city_unit_map.keys():
                    if known_city in city or city in known_city:
                        is_known = True
                        break
            if not is_known:
                unknown_cities.append(city)

    # 构建可序列化的报告
    report = {
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'cities': cities,
        'cities_str': summary.get('cities_str', ''),
        'unknown_cities': unknown_cities,
        'dates': {
            'earliest': date_objects.get('earliest').strftime('%Y-%m-%d') if date_objects.get('earliest') else None,
            'latest': date_objects.get('latest').strftime('%Y-%m-%d') if date_objects.get('latest') else None,
        },
        'amounts': {
            'jipiao': summary.get('jipiao_amount', 0),
            'huoche': summary.get('huoche_amount', 0),
            'tuigai': summary.get('tuigai_amount', 0),
            'didi': summary.get('didi_einvoice_amount', 0),
            'toll': summary.get('toll_amount', 0),
            'zhusu': summary.get('zhusu_amount', 0),
        },
        'counts': {
            'jipiao': summary.get('jipiao_count', 0),
            'huoche': summary.get('huoche_count', 0),
            'tuigai': summary.get('tuigai_count', 0),
            'didi': summary.get('didi_count', 0),
            'toll': summary.get('toll_count', 0),
        },
        'file_details': result.get('file_details', []),
        'warnings': result.get('warnings', []),
        'errors': result.get('errors', []),
        'needs_review': bool(result.get('warnings', [])) or len(unknown_cities) > 0,
        '_user_notes': '请复核以上数据，如有问题请修改后保存，然后执行 --write-data 阶段',
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    return report


def load_data_report(report_path='data_report.json'):
    """加载用户确认后的数据报告"""
    import json
    from datetime import datetime

    with open(report_path, 'r', encoding='utf-8') as f:
        report = json.load(f)

    # 转换日期字符串为datetime对象
    dates = report.get('dates', {})
    if dates.get('earliest'):
        dates['earliest'] = datetime.strptime(dates['earliest'], '%Y-%m-%d')
    if dates.get('latest'):
        dates['latest'] = datetime.strptime(dates['latest'], '%Y-%m-%d')

    return report


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(
        description='报销发票处理工具 - 三阶段执行',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
执行流程:
  阶段1 - 提取数据:  python3 invoice_processor.py --input-dir 发票 --extract-only
                      生成 data_report.json，供用户复核

  阶段2 - 写入数据:  python3 invoice_processor.py --write-data [--report data_report.json]
                      根据确认的数据写入Excel和Word

  阶段3 - 合并PDF:   python3 invoice_processor.py --merge-pdfs [--work-dir .]
                      转换并合并所有PDF

一键执行（无警告时自动完成，有警告时生成报告）:
  python3 invoice_processor.py --input-dir 发票 --output-excel biaoge.xlsx --auto
        """
    )

    # 阶段控制参数
    parser.add_argument('--extract-only', action='store_true', help='阶段1: 只提取数据，生成报告')
    parser.add_argument('--write-data', action='store_true', help='阶段2: 根据报告写入Excel和Word')
    parser.add_argument('--merge-pdfs', action='store_true', help='阶段3: PDF转换和合并')
    parser.add_argument('--auto', action='store_true', help='自动模式: 无警告时自动执行全部阶段')

    # 输入输出参数
    parser.add_argument('--input-dir', help='发票文件夹路径')
    parser.add_argument('--output-excel', help='输出Excel文件路径')
    parser.add_argument('--sheet', default='sheet1', help='工作表名称')
    parser.add_argument('--report', default='data_report.json', help='数据报告文件路径')
    parser.add_argument('--work-dir', default='.', help='工作目录（存放docx等文件）')
    parser.add_argument('--no-rename', action='store_true', help='不重命名文件')

    # 单位映射参数（阶段2使用）
    parser.add_argument('--city-units', help='城市到单位的映射，格式: "北京:总部,上海:分公司"')
    parser.add_argument('--config', default='config.yaml', help='配置文件路径（默认: config.yaml）')
    parser.add_argument('--force-mark', action='store_true', help='强制重新标记交易单号（忽略已有标记）')

    args = parser.parse_args()

    # 加载配置文件
    config = load_config(args.config)

    # 默认执行逻辑
    if not (args.extract_only or args.write_data or args.merge_pdfs or args.auto):
        if args.input_dir and args.output_excel:
            args.auto = True  # 有输入输出路径时启用自动模式
        else:
            parser.print_help()
            exit(1)

    # ==================== 阶段1: 提取数据 ====================
    if args.extract_only or args.auto:
        if not args.input_dir:
            print("❌ 请提供 --input-dir 参数")
            exit(1)

        # 每次重新提取前都使用干净发票副本，避免上次重命名、OFD转换、
        # PDF标记交易单号等处理结果影响本次识别。
        backup_dir = os.path.join(os.path.dirname(args.input_dir) or '.', 'bak')
        try:
            prepare_invoice_workspace(args.input_dir, backup_dir)
        except Exception as e:
            print(f"❌ 发票备份/恢复失败: {e}")
            exit(1)

        print("\n" + "="*60)
        print("【阶段1】提取发票数据")
        print("="*60)

        result = process_invoices(
            input_dir=args.input_dir,
            output_excel=None,  # 阶段1不写入Excel
            sheet_name=args.sheet,
            rename_files=not args.no_rename,
            force_write=True  # 阶段1不交互，只提取
        )

        # 生成数据报告
        report = generate_data_report(result, args.report, config)

        # 输出摘要
        print("\n" + "="*60)
        print("提取结果摘要")
        print("="*60)
        print(f"\n城市: {report['cities_str']}")
        print(f"机票: {report['counts']['jipiao']}张, ¥{report['amounts']['jipiao']}")
        print(f"日期: {report['dates']['earliest']} 至 {report['dates']['latest']}")
        print(f"住宿: ¥{report['amounts']['zhusu']}")
        print(f"滴滴: {report['counts']['didi']}张, ¥{report['amounts']['didi']}")
        print(f"高速费: {report['counts'].get('toll', 0)}张, ¥{report['amounts'].get('toll', 0)}")

        # 显示未知城市警告
        if report.get('unknown_cities'):
            print("\n" + "="*60)
            print("⚠️ 发现未知城市（未在config.yaml中配置）:")
            print("="*60)
            for city in report['unknown_cities']:
                print(f"  - {city}")           

        if report['warnings']:
            print("\n" + "="*60)
            print("⚠️ 发现以下问题需要复核:")
            print("="*60)
            for w in report['warnings']:
                print(f"  - {w}")
            print(f"\n✓ 数据报告已保存: {args.report}")
            print("  请复核并修改报告中的数据，然后执行:")
            print(f"  python3 {sys.argv[0]} --write-data --report {args.report}")

            if args.auto:
                print("\n(自动模式检测到警告，停止执行，请复核后再执行阶段2)")
                exit(0)  # 有警告时停止，等待用户复核
        else:
            print("\n✓ 数据提取完成，无警告")

        if args.extract_only:
            exit(0)

    # ==================== 阶段2: 写入数据 ====================
    if args.write_data or (args.auto and not args.extract_only):
        print("\n" + "="*60)
        print("【阶段2】写入数据到文档")
        print("="*60)

        # 加载数据报告
        if not os.path.exists(args.report):
            print(f"❌ 数据报告不存在: {args.report}")
            print("  请先执行阶段1: --extract-only")
            exit(1)

        report = load_data_report(args.report)

        # 解析用户提供的单位映射（命令行参数优先级最高）
        cli_city_unit_map = {}
        if args.city_units:
            for pair in args.city_units.split(','):
                if ':' in pair:
                    city, unit = pair.split(':', 1)
                    cli_city_unit_map[city.strip()] = unit.strip()

        # 获取最终的城市单位映射（命令行 > 配置文件）
        city_unit_map = get_city_unit_map(config, cli_city_unit_map)

        # 检查未知城市
        unknown_cities = report.get('unknown_cities', [])
        if unknown_cities:
            print("\n" + "="*60)
            print("⚠️ 发现以下未知城市需要处理:")
            print("="*60)
            for city in unknown_cities:
                print(f"  - {city}")

            # 检查是否已通过命令行参数提供单位
            still_unknown = [c for c in unknown_cities if c not in cli_city_unit_map]
            if still_unknown:
                print("\n这些城市未在config.yaml中配置，也未通过 --city-units 参数指定单位。")
                print("\n请在执行阶段2前选择以下方式之一:")
                print("  方式1: 修改 config.yaml 添加城市单位映射")
                print("  方式2: 使用 --city-units 参数指定单位，例如:")
                print(f"     --city-units \"{still_unknown[0]}:单位名称\"")
                print("\n或者，如果城市识别有误，请修改 data_report.json 中的 cities 字段后重新执行阶段2。")
                exit(1)
            else:
                print("\n✓ 已通过 --city-units 参数提供单位映射，继续执行...")

        # 写入Excel
        if args.output_excel:
            print(f"\n[步骤1] 写入Excel: {args.output_excel}")

            from openpyxl import load_workbook
            wb = load_workbook(args.output_excel)
            ws = wb[args.sheet]

            # 准备数据
            summary = {
                'cities_str': report['cities_str'],
                'huoche_count': report['counts']['huoche'],
                'huoche_amount': report['amounts']['huoche'],
                'jipiao_count': report['counts']['jipiao'],
                'jipiao_amount': report['amounts']['jipiao'],
                'tuigai_count': report['counts'].get('tuigai', 0),
                'tuigai_amount': report['amounts'].get('tuigai', 0),
                'didi_count': report['counts']['didi'],
                'didi_einvoice_amount': report['amounts']['didi'],
                'toll_count': report['counts'].get('toll', 0),
                'toll_amount': report['amounts'].get('toll', 0),
                'zhusu_amount': report['amounts']['zhusu'],
                'jipiao_earliest_date': report['dates']['earliest'],
                'jipiao_latest_date': report['dates']['latest'],
            }

            # 准备日期对象
            date_objects = {}
            if report['dates']['earliest']:
                date_objects['earliest'] = report['dates']['earliest']
            if report['dates']['latest']:
                date_objects['latest'] = report['dates']['latest']

            write_to_excel(args.output_excel, args.sheet, summary, date_objects)
            print("  ✓ Excel写入完成")

            # 复制日期到sheet2
            copy_dates_between_sheets(args.output_excel, args.sheet, 'sheet2')

            # 计算伙食补助费（定额：每天100元）
            daily_allowance = 100  # 每天100元伙食补助
            earliest_date = report['dates']['earliest']
            latest_date = report['dates']['latest']

            if earliest_date and latest_date:
                trip_days = (latest_date - earliest_date).days + 1  # 包含首尾两天
                meal_allowance = trip_days * daily_allowance
            else:
                trip_days = 0
                meal_allowance = 0

            # 更新 data_report.json，添加伙食补助信息
            report['trip_days'] = trip_days
            report['meal_allowance'] = meal_allowance
            report['meal_allowance_daily_rate'] = daily_allowance

            # 重新保存 data_report.json
            with open(args.report, 'w', encoding='utf-8') as f:
                # 转换日期为字符串以便JSON序列化
                report_for_save = report.copy()
                if isinstance(report_for_save['dates']['earliest'], datetime):
                    report_for_save['dates']['earliest'] = report_for_save['dates']['earliest'].strftime('%Y-%m-%d')
                if isinstance(report_for_save['dates']['latest'], datetime):
                    report_for_save['dates']['latest'] = report_for_save['dates']['latest'].strftime('%Y-%m-%d')
                json.dump(report_for_save, f, ensure_ascii=False, indent=2)
            print(f"  ✓ 已更新 {args.report}: 出差{trip_days}天, 伙食补助¥{meal_allowance}")

            # 计算总合计（所有费用 + 伙食补助）
            total_amount = (
                report['amounts']['jipiao'] +
                report['amounts']['huoche'] +
                report['amounts'].get('tuigai', 0) +
                report['amounts']['didi'] +
                report['amounts'].get('toll', 0) +
                report['amounts']['zhusu'] +
                meal_allowance
            )

            # 将合计转为中文大写，写入 E13
            try:
                chinese_total = amount_to_chinese(total_amount)
                wb = load_workbook(args.output_excel)
                ws = wb[args.sheet]
                ws['E13'] = chinese_total
                wb.save(args.output_excel)
                print(f"  ✓ 总合计¥{total_amount}已转为大写并写入E13: {chinese_total}")
            except Exception as e:
                print(f"  ⚠️ 写入E13大写金额失败: {e}")

        # 写入Word文档
        work_dir = args.work_dir
        docx_files = [f for f in os.listdir(work_dir) if f.endswith('.docx')]

        if docx_files and report['cities']:
            print(f"\n[步骤2] 写入Word文档")

            cities = report['cities']
            earliest_date = report['dates']['earliest']
            latest_date = report['dates']['latest']
            has_flight = report['counts']['jipiao'] > 0
            has_train = report['counts']['huoche'] > 0

            for docx_file in docx_files:
                docx_path = os.path.join(work_dir, docx_file)

                success, unknown_cities, _ = fill_docx_document(
                    docx_path,
                    cities,
                    earliest_date,
                    latest_date,
                    city_unit_map if city_unit_map else None,
                    has_flight,
                    has_train
                )

                if success:
                    print(f"  ✓ {docx_file} 写入完成")
                else:
                    print(f"  ⚠️ {docx_file} 写入可能不完整")

        print("\n✓ 阶段2完成: 数据已写入文档")

        if args.write_data:
            exit(0)

    # ==================== 阶段3: PDF转换和合并 ====================
    if args.merge_pdfs:
        print("\n" + "="*60)
        print("【阶段3】PDF转换和合并")
        print("="*60)

        work_dir = args.work_dir
        temp_pdfs = []

        # [步骤0] 微信支付账单处理 - 在转换前执行
        if args.input_dir:
            match_and_add_transaction_numbers(args.input_dir, work_dir, args.report, args.force_mark)
        else:
            print("\n[步骤0] 跳过微信支付账单处理（未提供 --input-dir）")

        # 确定Excel文件路径（优先使用命令行参数，否则自动查找）
        excel_path = args.output_excel
        if not excel_path:
            # 检查常见位置的biaoge.xlsx
            possible_paths = [
                os.path.join(work_dir, 'biaoge.xlsx'),
                'biaoge.xlsx',
                os.path.join(os.getcwd(), 'biaoge.xlsx'),
            ]
            for p in possible_paths:
                if os.path.exists(p):
                    excel_path = p
                    print(f"  自动找到Excel文件: {p}")
                    break

            if not excel_path:
                # 如果Excel不存在但PDF存在，跳过Excel转换
                biaoge_pdf_check = os.path.join(work_dir, 'biaoge.pdf')
                if os.path.exists(biaoge_pdf_check):
                    print(f"  ⚠️ 未找到Excel文件，但发现biaoge.pdf，将直接使用")
                    excel_path = None  # 标记为不转换
                else:
                    print("❌ 请提供 --output-excel 参数，或确保工作目录中有biaoge.xlsx或biaoge.pdf")
                    exit(1)

        # 1. 转换Excel为PDF
        print("\n[步骤1] 转换Excel为PDF")
        biaoge_pdf = os.path.join(work_dir, 'biaoge.pdf')
        if excel_path and convert_excel_to_pdf(excel_path, biaoge_pdf):
            temp_pdfs.append(biaoge_pdf)
            print("  ✓ Excel已转换")

        # 2. 转换Word为PDF（已经填写好的文档）
        print("\n[步骤2] 转换Word为PDF")
        docx_files = [f for f in os.listdir(work_dir) if f.endswith('.docx')]

        for docx_file in docx_files:
            docx_path = os.path.join(work_dir, docx_file)
            docx_pdf = os.path.join(work_dir, docx_file.replace('.docx', '.pdf'))

            if convert_docx_to_pdf(docx_path, docx_pdf):
                temp_pdfs.append(docx_pdf)
                print(f"  ✓ {docx_file} 已转换")

        # 3. 收集并合并所有PDF
        print("\n[步骤3] 收集并合并所有PDF文件")

        # 定义最终PDF列表
        final_pdf_list = []

        # biaoge.pdf放第一
        if os.path.exists(biaoge_pdf):
            final_pdf_list.append(biaoge_pdf)
            print(f"  ✓ 添加: {os.path.basename(biaoge_pdf)}")

        # shenpi.pdf放第二（如果存在）
        shenpi_pdf = os.path.join(work_dir, 'shenpi.pdf')
        if os.path.exists(shenpi_pdf):
            final_pdf_list.append(shenpi_pdf)
            print(f"  ✓ 添加: shenpi.pdf")

        # 收集发票PDF（按顺序）- 如果提供了input_dir
        if args.input_dir:
            invoice_pdfs = collect_pdfs_in_order(args.input_dir, work_dir)
            for pdf_file in invoice_pdfs:
                if pdf_file not in final_pdf_list:
                    final_pdf_list.append(pdf_file)
                    print(f"  ✓ 添加发票: {os.path.basename(pdf_file)}")
        else:
            print("  (未提供 --input-dir，跳过发票PDF收集)")

        # 4. 合并PDF
        if final_pdf_list:
            print(f"\n共收集 {len(final_pdf_list)} 个PDF文件")
            output_merge = os.path.join(work_dir, '汇总打印.pdf')
            merge_pdfs(final_pdf_list, output_merge)
        else:
            print("⚠️ 没有找到任何PDF文件可合并")

    print(f"\n处理完成!")
